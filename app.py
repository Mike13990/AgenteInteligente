import os
import re
import hashlib
import tempfile
from io import BytesIO
from datetime import datetime

import streamlit as st
import plotly.graph_objects as go

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from main import preparar_paquete_evaluacion, calificar_paquete
from db import guardar_examen_calificado


# ============================================================
# CONFIGURACIÓN DE PÁGINA
# ============================================================

st.set_page_config(
    page_title="EvaluaIA Neural | Grupo 8",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="collapsed"
)


# ============================================================
# CSS UI/UX
# ============================================================

st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;800;900&display=swap');

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
}

.stApp {
    background:
        radial-gradient(circle at top left, rgba(37,99,235,0.24), transparent 35%),
        radial-gradient(circle at top right, rgba(124,58,237,0.25), transparent 35%),
        radial-gradient(circle at bottom, rgba(14,165,233,0.14), transparent 45%),
        linear-gradient(135deg, #020617 0%, #0f172a 48%, #111827 100%);
    color: #e5e7eb;
}

.block-container {
    padding-top: 1.2rem;
    padding-bottom: 2.5rem;
    max-width: 1500px;
}

hr {
    border-color: rgba(148,163,184,0.18);
}

.hero {
    position: relative;
    overflow: hidden;
    border-radius: 34px;
    padding: 40px;
    margin-bottom: 24px;
    border: 1px solid rgba(125,211,252,0.30);
    background:
        linear-gradient(135deg, rgba(15,23,42,0.96), rgba(30,41,59,0.82)),
        radial-gradient(circle at 20% 20%, rgba(59,130,246,0.32), transparent 35%),
        radial-gradient(circle at 88% 10%, rgba(168,85,247,0.28), transparent 38%);
    box-shadow: 0 30px 90px rgba(0,0,0,0.48);
}

.hero::before {
    content: "";
    position: absolute;
    inset: -2px;
    background: linear-gradient(90deg, transparent, rgba(125,211,252,0.18), transparent);
    transform: translateX(-100%);
    animation: shine 5.5s infinite;
}

@keyframes shine {
    0% { transform: translateX(-100%); }
    55% { transform: translateX(120%); }
    100% { transform: translateX(120%); }
}

.hero-title {
    position: relative;
    font-size: 56px;
    line-height: 1;
    font-weight: 900;
    letter-spacing: -1.6px;
    background: linear-gradient(90deg, #93c5fd, #c4b5fd, #67e8f9);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
}

.hero-sub {
    position: relative;
    margin-top: 14px;
    max-width: 980px;
    color: #cbd5e1;
    font-size: 18px;
    line-height: 1.55;
}

.badge {
    display: inline-flex;
    align-items: center;
    gap: 6px;
    padding: 9px 14px;
    margin: 16px 8px 0 0;
    border-radius: 999px;
    color: #dbeafe;
    background: rgba(37,99,235,0.15);
    border: 1px solid rgba(96,165,250,0.35);
    font-size: 13px;
    font-weight: 800;
}

.glass {
    border-radius: 28px;
    padding: 22px;
    border: 1px solid rgba(148,163,184,0.20);
    background: rgba(15,23,42,0.62);
    box-shadow: 0 18px 55px rgba(0,0,0,0.28);
}

.card {
    border-radius: 26px;
    padding: 22px;
    min-height: 145px;
    border: 1px solid rgba(148,163,184,0.20);
    background: linear-gradient(180deg, rgba(15,23,42,0.88), rgba(2,6,23,0.72));
    box-shadow: 0 18px 48px rgba(0,0,0,0.30);
}

.card:hover {
    transform: translateY(-2px);
    border-color: rgba(125,211,252,0.45);
    box-shadow: 0 24px 60px rgba(0,0,0,0.40);
    transition: 0.2s ease;
}

.card-label {
    color: #93c5fd;
    font-size: 13px;
    font-weight: 900;
    text-transform: uppercase;
    letter-spacing: .10em;
}

.card-value {
    margin-top: 8px;
    color: #f8fafc;
    font-size: 30px;
    font-weight: 900;
}

.card-desc {
    margin-top: 8px;
    color: #94a3b8;
    font-size: 14px;
    line-height: 1.45;
}

.upload-box {
    border-radius: 24px;
    padding: 22px;
    border: 1px dashed rgba(125,211,252,0.38);
    background: linear-gradient(180deg, rgba(15,23,42,0.72), rgba(30,41,59,0.38));
    margin-bottom: 14px;
}

.section-title {
    font-size: 28px;
    font-weight: 900;
    color: #f8fafc;
    margin: 12px 0 6px 0;
}

.section-sub {
    color: #94a3b8;
    margin-bottom: 16px;
}

.status-chip {
    display: inline-block;
    padding: 8px 12px;
    border-radius: 999px;
    background: rgba(59,130,246,0.16);
    border: 1px solid rgba(96,165,250,0.32);
    color: #dbeafe;
    font-weight: 800;
    font-size: 13px;
}

.difficulty {
    border-radius: 24px;
    padding: 20px;
    margin-top: 12px;
    border: 1px solid rgba(148,163,184,0.22);
    font-weight: 800;
}

.diff-green { background: rgba(34,197,94,0.15); color: #bbf7d0; border-color: rgba(34,197,94,0.35); }
.diff-yellow { background: rgba(234,179,8,0.15); color: #fef3c7; border-color: rgba(234,179,8,0.35); }
.diff-blue { background: rgba(59,130,246,0.15); color: #bfdbfe; border-color: rgba(59,130,246,0.35); }
.diff-orange { background: rgba(249,115,22,0.15); color: #fed7aa; border-color: rgba(249,115,22,0.35); }
.diff-red { background: rgba(239,68,68,0.15); color: #fecaca; border-color: rgba(239,68,68,0.35); }

.stButton > button {
    width: 100%;
    border-radius: 18px;
    padding: 15px 18px;
    font-size: 17px;
    font-weight: 900;
    color: white;
    border: 1px solid rgba(125,211,252,0.45);
    background: linear-gradient(90deg, #2563eb, #7c3aed);
    box-shadow: 0 15px 36px rgba(37,99,235,0.26);
}

.stButton > button:hover {
    transform: translateY(-1px);
    border-color: rgba(125,211,252,0.85);
    box-shadow: 0 20px 50px rgba(124,58,237,0.34);
}

.stDownloadButton > button {
    width: 100%;
    border-radius: 18px;
    padding: 14px 18px;
    font-weight: 900;
    color: white;
    background: linear-gradient(90deg, #0f766e, #2563eb);
    border: 1px solid rgba(125,211,252,0.35);
}

div[data-testid="stFileUploader"] section {
    background: rgba(2,6,23,0.45);
    border: 1px dashed rgba(125,211,252,0.35);
    border-radius: 18px;
}

.result-shell {
    border-radius: 30px;
    padding: 26px;
    border: 1px solid rgba(148,163,184,0.22);
    background: linear-gradient(180deg, rgba(15,23,42,0.88), rgba(2,6,23,0.72));
    box-shadow: 0 22px 72px rgba(0,0,0,0.38);
}

.question-card {
    border-radius: 22px;
    padding: 20px;
    border: 1px solid rgba(148,163,184,0.16);
    background: rgba(15,23,42,0.68);
}

.loading-card {
    border-radius: 24px;
    padding: 18px 20px;
    border: 1px solid rgba(96,165,250,0.28);
    background: rgba(15,23,42,0.76);
    color: #dbeafe;
    font-weight: 800;
    box-shadow: 0 0 35px rgba(59,130,246,0.16);
}

.small {
    color: #94a3b8;
    font-size: 13px;
}

.footer-note {
    color: #64748b;
    font-size: 13px;
    text-align: center;
    padding-top: 18px;
}
</style>
""",
    unsafe_allow_html=True
)


# ============================================================
# SESSION STATE
# ============================================================

defaults = {
    "paquete_cache": None,
    "ultimo_hash": None,
    "ultimo_resultado": None,
    "ultima_dificultad": 5,
    "mostrar_resultado": False,
    "contextos_bytes": [],
    "imagenes_bytes": [],
    "ultimo_resumen": {},
    "ultimo_guardado_id": None,
    "datos_reporte": {},
    "metadatos_aplicados": False,
}

for key, value in defaults.items():
    if key not in st.session_state:
        st.session_state[key] = value

# Compatibilidad si venías usando imagen_bytes en una versión anterior
if "imagen_bytes" in st.session_state and st.session_state.imagen_bytes and not st.session_state.imagenes_bytes:
    st.session_state.imagenes_bytes = [st.session_state.imagen_bytes]


# ============================================================
# FUNCIONES GENERALES
# ============================================================

def hash_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def limpiar_markdown(texto: str) -> str:
    return (
        str(texto)
        .replace("```markdown", "")
        .replace("```", "")
        .strip()
    )


def limpiar_para_archivo(texto: str) -> str:
    texto = texto or "SinDato"
    texto = texto.strip()
    texto = re.sub(r"[^\w\sáéíóúÁÉÍÓÚñÑ-]", "", texto)
    texto = re.sub(r"\s+", "_", texto)
    return texto[:60] or "SinDato"


def nombre_archivo_reporte(extension: str, examen_id=None, estudiante=None, curso=None) -> str:
    fecha = datetime.now().strftime("%Y-%m-%d")
    id_txt = f"ID{examen_id}" if examen_id else "SinID"
    estudiante_txt = limpiar_para_archivo(estudiante)
    curso_txt = limpiar_para_archivo(curso)
    return f"Reporte_{id_txt}_{estudiante_txt}_{curso_txt}_{fecha}.{extension}"


def guardar_archivo_temporal(temp_dir, nombre, data):
    nombre_seguro = limpiar_para_archivo(os.path.splitext(nombre)[0]) + os.path.splitext(nombre)[1]
    ruta = os.path.join(temp_dir, nombre_seguro)
    with open(ruta, "wb") as f:
        f.write(data)
    return ruta


def extraer_numero(patron, texto, default=None):
    match = re.search(patron, texto, flags=re.IGNORECASE)
    if not match:
        return default

    try:
        return float(match.group(1).replace(",", "."))
    except Exception:
        return default


def extraer_entero(patron, texto, default=0):
    match = re.search(patron, texto, flags=re.IGNORECASE)
    if not match:
        return default

    try:
        return int(float(match.group(1).replace(",", ".")))
    except Exception:
        return default


def extraer_texto_seccion(titulo, texto):
    patron = rf"## {re.escape(titulo)}\s*(.*?)(?=\n## |\Z)"
    match = re.search(patron, texto, flags=re.IGNORECASE | re.DOTALL)
    if not match:
        return None
    return match.group(1).strip()


def quitar_markdown_para_reporte(texto: str) -> str:
    texto = limpiar_markdown(texto)
    texto = re.sub(r"#{1,6}\s*", "", texto)
    texto = texto.replace("**", "")
    texto = texto.replace("*", "")
    texto = texto.replace("`", "")
    return texto.strip()


def extraer_resumen(resultado: str) -> dict:
    texto = limpiar_markdown(resultado)

    obtenido = extraer_numero(
        r"Punteo obtenido:\*\*\s*([0-9]+(?:[.,][0-9]+)?)",
        texto,
        None
    )

    total = extraer_numero(
        r"Punteo obtenido:\*\*\s*[0-9]+(?:[.,][0-9]+)?\s*/\s*([0-9]+(?:[.,][0-9]+)?)",
        texto,
        None
    )

    nota = extraer_numero(
        r"Nota final en escala de 100:\*\*\s*([0-9]+(?:[.,][0-9]+)?)",
        texto,
        None
    )

    nivel = extraer_entero(
        r"Nivel de dificultad aplicado:\*\*\s*([0-9]+)",
        texto,
        st.session_state.ultima_dificultad
    )

    buenas = extraer_entero(r"Buenas:\*\*\s*([0-9]+)", texto, 0)
    parciales = extraer_entero(r"Parciales:\*\*\s*([0-9]+)", texto, 0)
    incorrectas = extraer_entero(r"Incorrectas:\*\*\s*([0-9]+)", texto, 0)

    return {
        "obtenido": obtenido,
        "total": total,
        "nota": nota,
        "nivel": nivel,
        "buenas": buenas,
        "parciales": parciales,
        "incorrectas": incorrectas,
    }


def dificultad_info(nivel: int):
    if nivel <= 2:
        return "Muy flexible", "Prioriza acercamiento e intención de respuesta.", "diff-green", "🟢"
    if nivel <= 4:
        return "Flexible", "Acepta respuestas breves si capturan la idea central.", "diff-yellow", "🟡"
    if nivel <= 6:
        return "Estándar académico", "Balance entre justicia, comprensión y precisión.", "diff-blue", "🔵"
    if nivel <= 8:
        return "Estricto", "Penaliza omisiones, vaguedad y falta de desarrollo.", "diff-orange", "🟠"
    return "Experto riguroso", "Exige precisión técnica, completitud y profundidad.", "diff-red", "🔴"


def card(label, value, desc, icon="🧠"):
    st.markdown(
        f"""
        <div class="card">
            <div class="card-label">{icon} {label}</div>
            <div class="card-value">{value}</div>
            <div class="card-desc">{desc}</div>
        </div>
        """,
        unsafe_allow_html=True
    )


def metric_card(label, value, desc):
    st.markdown(
        f"""
        <div class="card">
            <div class="card-label">{label}</div>
            <div class="card-value">{value}</div>
            <div class="card-desc">{desc}</div>
        </div>
        """,
        unsafe_allow_html=True
    )


def crear_gauge(nota):
    nota = nota or 0

    fig = go.Figure(
        go.Indicator(
            mode="gauge+number",
            value=nota,
            number={"suffix": "/100", "font": {"size": 36, "color": "#f8fafc"}},
            gauge={
                "axis": {"range": [0, 100], "tickcolor": "#94a3b8"},
                "bar": {"color": "#60a5fa"},
                "bgcolor": "rgba(15,23,42,0.55)",
                "borderwidth": 1,
                "bordercolor": "rgba(148,163,184,0.35)",
                "steps": [
                    {"range": [0, 60], "color": "rgba(239,68,68,0.25)"},
                    {"range": [60, 80], "color": "rgba(245,158,11,0.25)"},
                    {"range": [80, 100], "color": "rgba(34,197,94,0.25)"},
                ],
            },
        )
    )

    fig.update_layout(
        height=320,
        margin=dict(l=10, r=10, t=20, b=10),
        paper_bgcolor="rgba(0,0,0,0)",
        font=dict(color="#e5e7eb"),
    )

    return fig


def crear_donut(buenas, parciales, incorrectas):
    fig = go.Figure(
        data=[
            go.Pie(
                labels=["Buenas", "Parciales", "Incorrectas"],
                values=[buenas, parciales, incorrectas],
                hole=0.62,
                textinfo="label+value",
                marker=dict(colors=["#22c55e", "#f59e0b", "#ef4444"]),
            )
        ]
    )

    fig.update_layout(
        height=320,
        margin=dict(l=10, r=10, t=20, b=10),
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        font=dict(color="#e5e7eb", size=14),
        showlegend=False,
    )

    return fig


def separar_preguntas(resultado: str):
    texto = limpiar_markdown(resultado)
    partes = re.split(r"(?=### Pregunta global\s+\d+)", texto)
    intro = partes[0].strip() if partes else ""
    preguntas = [p.strip() for p in partes[1:] if p.strip()]
    return intro, preguntas


def estado_pregunta(bloque: str):
    if "Estado:** Correcta" in bloque or "Estado: Correcta" in bloque:
        return "✅"
    if "Estado:** Parcial" in bloque or "Estado: Parcial" in bloque:
        return "🟡"
    if "Estado:** Incorrecta" in bloque or "Estado: Incorrecta" in bloque:
        return "❌"
    return "🧠"


def safe_image(image_bytes: bytes, **kwargs):
    try:
        st.image(image_bytes, width="stretch", **kwargs)
    except TypeError:
        st.image(image_bytes, use_container_width=True, **kwargs)


def aplicar_metadatos_detectados(paquete: dict):
    """
    Llena datos_reporte una sola vez con lo detectado por el OCR/ICR.
    El docente siempre puede editar después.
    """
    metadatos = paquete.get("metadatos") or {}

    if not metadatos:
        return

    actuales = st.session_state.datos_reporte or {}

    nuevos = {
        "examen_id": st.session_state.ultimo_guardado_id,
        "estudiante": actuales.get("estudiante") or metadatos.get("estudiante", ""),
        "codigo": actuales.get("codigo") or metadatos.get("codigo", ""),
        "curso": actuales.get("curso") or metadatos.get("curso", ""),
        "docente": actuales.get("docente") or metadatos.get("docente", ""),
        "titulo": actuales.get("titulo") or metadatos.get("titulo", ""),
        "serie": actuales.get("serie") or metadatos.get("serie", ""),
        "fecha_examen": actuales.get("fecha_examen") or metadatos.get("fecha_examen", ""),
    }

    st.session_state.datos_reporte = nuevos
    st.session_state.metadatos_aplicados = True


def resumen_archivos_examen(imagenes: list) -> tuple[str, str]:
    if not imagenes:
        return None, None

    nombres = [img.get("name", "imagen") for img in imagenes]
    hashes = [hash_bytes(img.get("bytes", b"")) for img in imagenes]

    archivo_nombre = " | ".join(nombres)
    archivo_hash = hashlib.sha256("".join(hashes).encode()).hexdigest()

    return archivo_nombre, archivo_hash


# ============================================================
# GENERADORES DE PDF Y WORD
# ============================================================

def marca_agua_pdf(canvas, doc):
    canvas.saveState()

    width, height = letter

    canvas.setFillColor(colors.Color(0.20, 0.36, 0.65, alpha=0.10))
    canvas.setFont("Helvetica-Bold", 46)

    canvas.translate(width / 2, height / 2)
    canvas.rotate(35)
    canvas.drawCentredString(0, 0.45 * inch, "GRUPO 8")
    canvas.setFont("Helvetica-Bold", 24)
    canvas.drawCentredString(0, 0, "EvaluaIA Neural")
    canvas.setFont("Helvetica", 16)
    canvas.drawCentredString(0, -0.35 * inch, "Sistema Inteligente de Calificación")

    canvas.restoreState()

    canvas.saveState()
    canvas.setFillColor(colors.HexColor("#64748B"))
    canvas.setFont("Helvetica", 8)
    canvas.drawString(0.7 * inch, 0.45 * inch, "Generado por Grupo 8 · EvaluaIA Neural")
    canvas.drawRightString(width - 0.7 * inch, 0.45 * inch, f"Página {doc.page}")
    canvas.restoreState()


def generar_pdf_reporte(datos: dict, resumen: dict, informe_markdown: str) -> bytes:
    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.7 * inch,
        leftMargin=0.7 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=22,
        leading=26,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#1E3A8A"),
        spaceAfter=12,
    )

    subtitle_style = ParagraphStyle(
        "CustomSubtitle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#475569"),
        spaceAfter=18,
    )

    h_style = ParagraphStyle(
        "CustomH",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#1E40AF"),
        spaceBefore=10,
        spaceAfter=8,
    )

    body_style = ParagraphStyle(
        "CustomBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#111827"),
        alignment=TA_LEFT,
    )

    story = []

    story.append(Paragraph("Reporte de Calificación Inteligente", title_style))
    story.append(Paragraph("GRUPO 8 · EvaluaIA Neural · Sistema Inteligente Multimodal", subtitle_style))

    info = [
        ["ID historial", str(datos.get("examen_id") or "Sin guardar")],
        ["Estudiante", datos.get("estudiante") or "No especificado"],
        ["Código / carné", datos.get("codigo") or "No especificado"],
        ["Curso", datos.get("curso") or "No especificado"],
        ["Docente", datos.get("docente") or "No especificado"],
        ["Título del examen", datos.get("titulo") or "No especificado"],
        ["Serie", datos.get("serie") or "No especificado"],
        ["Fecha del examen", datos.get("fecha_examen") or "No especificada"],
        ["Fecha de descarga", datetime.now().strftime("%Y-%m-%d %H:%M")],
    ]

    tabla_info = Table(info, colWidths=[1.75 * inch, 4.85 * inch])
    tabla_info.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#DBEAFE")),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#1E3A8A")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (1, 0), (1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
                ("PADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )

    story.append(Paragraph("Datos generales", h_style))
    story.append(tabla_info)
    story.append(Spacer(1, 14))

    obtenido = resumen.get("obtenido")
    total = resumen.get("total")
    nota = resumen.get("nota")
    nivel = resumen.get("nivel")
    buenas = resumen.get("buenas", 0)
    parciales = resumen.get("parciales", 0)
    incorrectas = resumen.get("incorrectas", 0)

    tabla_resumen = [
        ["Punteo", "Nota /100", "Nivel", "Buenas", "Parciales", "Incorrectas"],
        [
            f"{obtenido:g}/{total:g}" if obtenido is not None and total else "—",
            f"{nota:g}/100" if nota is not None else "—",
            str(nivel or "—"),
            str(buenas),
            str(parciales),
            str(incorrectas),
        ],
    ]

    t_resumen = Table(tabla_resumen, colWidths=[1.1 * inch] * 6)
    t_resumen.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor("#EFF6FF")),
                ("FONTNAME", (0, 1), (-1, 1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
                ("PADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )

    story.append(Paragraph("Resumen de calificación", h_style))
    story.append(t_resumen)
    story.append(Spacer(1, 14))

    story.append(Paragraph("Informe completo generado por IA", h_style))

    texto_limpio = quitar_markdown_para_reporte(informe_markdown)
    parrafos = texto_limpio.split("\n")

    for parrafo in parrafos:
        parrafo = parrafo.strip()
        if not parrafo:
            story.append(Spacer(1, 5))
        else:
            parrafo = parrafo.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(parrafo, body_style))

    doc.build(story, onFirstPage=marca_agua_pdf, onLaterPages=marca_agua_pdf)

    pdf = buffer.getvalue()
    buffer.close()
    return pdf


def generar_word_reporte(datos: dict, resumen: dict, informe_markdown: str) -> bytes:
    buffer = BytesIO()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    header = section.header
    h = header.paragraphs[0]
    h.text = "GRUPO 8 · EvaluaIA Neural · Sistema Inteligente de Calificación"
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer
    f = footer.paragraphs[0]
    f.text = "Generado por Grupo 8 · EvaluaIA Neural"
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER

    watermark = doc.add_paragraph()
    watermark.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_wm = watermark.add_run("GRUPO 8 · EvaluaIA Neural")
    run_wm.bold = True
    run_wm.font.size = Pt(28)
    run_wm.font.color.rgb = RGBColor(180, 190, 210)

    title = doc.add_heading("Reporte de Calificación Inteligente", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("GRUPO 8 · EvaluaIA Neural · Sistema Inteligente Multimodal")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Datos generales", level=1)

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    info = [
        ("ID historial", str(datos.get("examen_id") or "Sin guardar")),
        ("Estudiante", datos.get("estudiante") or "No especificado"),
        ("Código / carné", datos.get("codigo") or "No especificado"),
        ("Curso", datos.get("curso") or "No especificado"),
        ("Docente", datos.get("docente") or "No especificado"),
        ("Título del examen", datos.get("titulo") or "No especificado"),
        ("Serie", datos.get("serie") or "No especificado"),
        ("Fecha del examen", datos.get("fecha_examen") or "No especificada"),
        ("Fecha de descarga", datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]

    for etiqueta, valor in info:
        row = table.add_row().cells
        row[0].text = etiqueta
        row[1].text = valor

    doc.add_heading("Resumen de calificación", level=1)

    obtenido = resumen.get("obtenido")
    total = resumen.get("total")
    nota = resumen.get("nota")
    nivel = resumen.get("nivel")
    buenas = resumen.get("buenas", 0)
    parciales = resumen.get("parciales", 0)
    incorrectas = resumen.get("incorrectas", 0)

    table2 = doc.add_table(rows=2, cols=6)
    table2.style = "Table Grid"

    headers = ["Punteo", "Nota /100", "Nivel", "Buenas", "Parciales", "Incorrectas"]
    values = [
        f"{obtenido:g}/{total:g}" if obtenido is not None and total else "—",
        f"{nota:g}/100" if nota is not None else "—",
        str(nivel or "—"),
        str(buenas),
        str(parciales),
        str(incorrectas),
    ]

    for i, header_text in enumerate(headers):
        table2.rows[0].cells[i].text = header_text
        table2.rows[1].cells[i].text = values[i]

    doc.add_heading("Informe completo generado por IA", level=1)

    texto_limpio = quitar_markdown_para_reporte(informe_markdown)

    for parrafo in texto_limpio.split("\n"):
        parrafo = parrafo.strip()
        if parrafo:
            doc.add_paragraph(parrafo)

    doc.save(buffer)
    data = buffer.getvalue()
    buffer.close()
    return data


# ============================================================
# MODALES
# ============================================================

if hasattr(st, "dialog"):

    @st.dialog("🧠 Arquitectura inteligente")
    def modal_arquitectura():
        st.markdown(
            """
### Pipeline de EvaluaIA Neural

**1. OCR/ICR Vision**  
Lee una o varias imágenes del examen, detecta encabezado, preguntas, respuestas y rúbrica.

**2. RAG Contextual**  
Busca fragmentos relevantes en uno o varios documentos del profesor.

**3. LLM Calificador**  
Compara la respuesta del estudiante contra el material de referencia.

**4. Dificultad Adaptativa**  
Ajusta la exigencia de 1 a 10.

**5. Metadatos automáticos**  
Intenta detectar estudiante, curso, docente, serie y fecha.

**6. Reportes PDF / Word**  
Genera documentos descargables con marca de agua del Grupo 8.

**7. Historial PostgreSQL**  
Guarda resultados para consultarlos después.
"""
        )

    @st.dialog("📘 Guía rápida")
    def modal_guia():
        st.markdown(
            """
### Cómo usar el sistema

1. Sube uno o varios **materiales de referencia**.
2. Sube una o varias **imágenes del mismo examen**.
3. Selecciona el nivel de dificultad.
4. Presiona **Ejecutar análisis inteligente**.
5. Revisa el dashboard.
6. Verifica o edita los datos detectados.
7. Descarga el reporte en **PDF** o **Word**.
8. Guarda la calificación en historial.
"""
        )


# ============================================================
# HERO
# ============================================================

st.markdown(
    """
    <div class="hero">
        <div class="hero-title">EvaluaIA Neural</div>
        <div class="hero-sub">
            Sistema inteligente multimodal para la calificación autónoma de exámenes.
            Integra visión artificial, OCR/ICR, RAG, LLM, reportes académicos,
            detección automática de metadatos y almacenamiento histórico.
        </div>
        <div>
            <span class="badge">👁️ OCR/ICR Vision</span>
            <span class="badge">📚 RAG Contextual</span>
            <span class="badge">🧠 LLM Calificador</span>
            <span class="badge">🖼️ Multi-imagen</span>
            <span class="badge">🎯 Dificultad 1–10</span>
            <span class="badge">📄 PDF / Word</span>
            <span class="badge">🗄️ Historial PostgreSQL</span>
        </div>
    </div>
    """,
    unsafe_allow_html=True
)


# ============================================================
# BOTONES SUPERIORES
# ============================================================

b1, b2, b3, b4 = st.columns([1, 1, 1, 2])

with b1:
    if st.button("🧠 Arquitectura IA"):
        if hasattr(st, "dialog"):
            modal_arquitectura()
        else:
            st.info("OCR/ICR → RAG → LLM → Reporte → Historial")

with b2:
    if st.button("📘 Guía rápida"):
        if hasattr(st, "dialog"):
            modal_guia()
        else:
            st.info("Sube material, sube examen, elige dificultad, califica, descarga y guarda.")

with b3:
    if st.button("📚 Ver historial"):
        st.switch_page("pages/historial.py")

with b4:
    st.markdown(
        """
        <div class="glass">
            <span class="status-chip">Grupo 8</span>
            <span class="small"> · Reportes con marca de agua y metadatos detectados.</span>
        </div>
        """,
        unsafe_allow_html=True
    )


# ============================================================
# TARJETAS
# ============================================================

st.divider()

c1, c2, c3, c4 = st.columns(4)

with c1:
    card("Materiales", str(len(st.session_state.contextos_bytes)), "Documentos usados como base de conocimiento.", "📚")

with c2:
    card(
        "Imágenes examen",
        str(len(st.session_state.imagenes_bytes)) if st.session_state.imagenes_bytes else "Pendiente",
        "Una o varias hojas del mismo examen.",
        "📝"
    )

with c3:
    card("Motor IA", "OCR + RAG + LLM", "Procesamiento visual y evaluación contextual.", "⚡")

with c4:
    card("Reportes", "PDF / Word", "Descarga con nombre, curso, ID y marca de agua.", "📄")


# ============================================================
# CARGA DE ARCHIVOS
# ============================================================

st.markdown('<div class="section-title">📥 Carga de archivos</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="section-sub">Sube los materiales del profesor y una o varias imágenes del mismo examen.</div>',
    unsafe_allow_html=True
)

col1, col2 = st.columns(2)

with col1:
    st.markdown(
        """
        <div class="upload-box">
            <h3>📚 Materiales de referencia</h3>
            <p style="color:#94a3b8;">PDF, DOCX, TXT o MD. Puedes subir varios documentos.</p>
        </div>
        """,
        unsafe_allow_html=True
    )

    archivos_contexto = st.file_uploader(
        "Materiales",
        type=["pdf", "txt", "md", "docx"],
        accept_multiple_files=True,
        key="uploader_contextos",
        label_visibility="collapsed"
    )

with col2:
    st.markdown(
        """
        <div class="upload-box">
            <h3>📝 Examen del estudiante</h3>
            <p style="color:#94a3b8;">PNG, JPG o JPEG. Puedes subir varias hojas/imágenes del mismo examen.</p>
        </div>
        """,
        unsafe_allow_html=True
    )

    imagenes_examen = st.file_uploader(
        "Examen",
        type=["png", "jpg", "jpeg"],
        accept_multiple_files=True,
        key="uploader_imagenes",
        label_visibility="collapsed"
    )


if archivos_contexto:
    nuevos_contextos = [
        {"name": archivo.name, "bytes": archivo.getvalue()}
        for archivo in archivos_contexto
    ]
    st.session_state.contextos_bytes = nuevos_contextos

if imagenes_examen:
    nuevas_imagenes = [
        {"name": imagen.name, "bytes": imagen.getvalue()}
        for imagen in imagenes_examen
    ]
    st.session_state.imagenes_bytes = nuevas_imagenes


# ============================================================
# PREVISUALIZACIÓN
# ============================================================

p1, p2 = st.columns(2)

with p1:
    st.markdown("### 📚 Material cargado")
    if st.session_state.contextos_bytes:
        for archivo in st.session_state.contextos_bytes:
            size_kb = len(archivo["bytes"]) / 1024
            st.success(f"{archivo['name']} · {size_kb:.1f} KB")
    else:
        st.info("Aún no has subido material de referencia.")

with p2:
    st.markdown("### 🖼️ Vista previa del examen")
    if st.session_state.imagenes_bytes:
        st.success(f"{len(st.session_state.imagenes_bytes)} imagen(es) cargada(s).")
        tabs = st.tabs([f"Hoja {i}" for i in range(1, len(st.session_state.imagenes_bytes) + 1)])
        for tab, imagen in zip(tabs, st.session_state.imagenes_bytes):
            with tab:
                st.caption(imagen["name"])
                safe_image(imagen["bytes"])
    else:
        st.info("Aún no has subido imágenes del examen.")


# ============================================================
# CONFIGURACIÓN DE CALIFICACIÓN
# ============================================================

st.divider()
st.markdown('<div class="section-title">⚙️ Configuración de calificación</div>', unsafe_allow_html=True)

nivel_dificultad = st.slider(
    "Nivel de dificultad",
    min_value=1,
    max_value=10,
    value=st.session_state.ultima_dificultad,
    step=1
)

st.session_state.ultima_dificultad = nivel_dificultad

modo, descripcion, clase, icono = dificultad_info(nivel_dificultad)

st.markdown(
    f"""
    <div class="difficulty {clase}">
        {icono} <strong>{modo}</strong> · Nivel {nivel_dificultad}/10
        <br>
        <span style="font-weight:600;">{descripcion}</span>
    </div>
    """,
    unsafe_allow_html=True
)

mostrar_extraccion = st.checkbox("🔎 Mostrar extracción OCR/ICR después del análisis", value=False)


# ============================================================
# BOTONES DE ACCIÓN
# ============================================================

st.divider()

a1, a2 = st.columns([4, 1])

with a1:
    iniciar = st.button("🚀 Ejecutar análisis inteligente")

with a2:
    limpiar = st.button("🗑️ Reiniciar")


if limpiar:
    st.session_state.paquete_cache = None
    st.session_state.ultimo_hash = None
    st.session_state.ultimo_resultado = None
    st.session_state.mostrar_resultado = False
    st.session_state.contextos_bytes = []
    st.session_state.imagenes_bytes = []
    st.session_state.ultimo_resumen = {}
    st.session_state.ultimo_guardado_id = None
    st.session_state.datos_reporte = {}
    st.session_state.metadatos_aplicados = False
    if "imagen_bytes" in st.session_state:
        st.session_state.imagen_bytes = None
    st.rerun()


# ============================================================
# PROCESO DE ANÁLISIS
# ============================================================

if iniciar:
    if not st.session_state.contextos_bytes:
        st.error("📚 Falta el material de referencia. Sube al menos un PDF, Word, TXT o MD.")
        st.stop()

    if not st.session_state.imagenes_bytes:
        st.error("📝 Falta el examen. Sube al menos una imagen del examen.")
        st.stop()

    if len(st.session_state.imagenes_bytes) > 8:
        st.warning("⚠️ Subiste más de 8 imágenes. Puede funcionar, pero tardará más y consumirá más tokens.")

    try:
        hash_partes = []

        for imagen in st.session_state.imagenes_bytes:
            hash_partes.append(hash_bytes(imagen["bytes"]))

        for archivo in st.session_state.contextos_bytes:
            hash_partes.append(hash_bytes(archivo["bytes"]))

        hash_partes.append(str(nivel_dificultad))
        hash_actual = hashlib.sha256("".join(hash_partes).encode()).hexdigest()

        reutilizar_cache = (
            st.session_state.paquete_cache is not None
            and st.session_state.ultimo_hash == hash_actual
        )

        progress = st.progress(0)
        estado = st.empty()

        with tempfile.TemporaryDirectory() as temp_dir:
            estado.markdown(
                '<div class="loading-card">🧩 Preparando archivos temporales...</div>',
                unsafe_allow_html=True
            )
            progress.progress(10)

            rutas_imagenes = []
            for i, imagen in enumerate(st.session_state.imagenes_bytes, start=1):
                ruta = guardar_archivo_temporal(
                    temp_dir,
                    f"hoja_{i}_{imagen['name']}",
                    imagen["bytes"]
                )
                rutas_imagenes.append(ruta)

            rutas_contexto = []
            for archivo in st.session_state.contextos_bytes:
                ruta = guardar_archivo_temporal(
                    temp_dir,
                    archivo["name"],
                    archivo["bytes"]
                )
                rutas_contexto.append(ruta)

            if reutilizar_cache:
                estado.markdown(
                    '<div class="loading-card">♻️ Reutilizando OCR/ICR y RAG procesados...</div>',
                    unsafe_allow_html=True
                )
                progress.progress(55)
                paquete = st.session_state.paquete_cache

            else:
                estado.markdown(
                    '<div class="loading-card">👁️ Leyendo imágenes con OCR/ICR y detectando encabezado, rúbrica y respuestas...</div>',
                    unsafe_allow_html=True
                )
                progress.progress(25)

                with st.spinner("Procesando visión y contexto..."):
                    paquete = preparar_paquete_evaluacion(
                        rutas_imagenes,
                        rutas_contexto
                    )

                progress.progress(68)

                if not isinstance(paquete, dict):
                    st.error("El motor principal no devolvió un paquete válido.")
                    st.stop()

                if "error" in paquete:
                    st.error(paquete["error"])
                    st.stop()

                st.session_state.paquete_cache = paquete
                st.session_state.ultimo_hash = hash_actual

            aplicar_metadatos_detectados(paquete)

            if mostrar_extraccion:
                with st.expander("📄 Texto extraído del examen", expanded=True):
                    st.text(paquete.get("texto_examen", ""))

                with st.expander("🧾 Metadatos detectados", expanded=True):
                    st.json(paquete.get("metadatos", {}))

                with st.expander("🎯 Rúbrica detectada", expanded=False):
                    st.json(paquete.get("rubrica", {}))

            estado.markdown(
                '<div class="loading-card">📚 Comparando respuestas contra el material de referencia...</div>',
                unsafe_allow_html=True
            )
            progress.progress(78)

            with st.spinner("Calificando con IA..."):
                resultado = calificar_paquete(
                    paquete,
                    nivel_dificultad
                )

            estado.markdown(
                '<div class="loading-card">📑 Generando informe académico y dashboard...</div>',
                unsafe_allow_html=True
            )
            progress.progress(92)

            resultado = limpiar_markdown(resultado)

            if not resultado or resultado.startswith("ERROR:"):
                st.error(resultado or "La IA terminó, pero no devolvió un reporte válido.")
                st.stop()

            st.session_state.ultimo_resultado = resultado
            st.session_state.ultimo_resumen = extraer_resumen(resultado)
            st.session_state.mostrar_resultado = True
            st.session_state.ultimo_guardado_id = None
            st.session_state.datos_reporte["examen_id"] = None

            progress.progress(100)
            estado.markdown(
                '<div class="loading-card">✅ Evaluación completada.</div>',
                unsafe_allow_html=True
            )

            st.rerun()

    except Exception as e:
        st.error("❌ Error general durante el análisis.")
        st.exception(e)


# ============================================================
# RESULTADOS
# ============================================================

if st.session_state.mostrar_resultado and st.session_state.ultimo_resultado:
    resultado_limpio = limpiar_markdown(st.session_state.ultimo_resultado)
    resumen = st.session_state.ultimo_resumen or extraer_resumen(resultado_limpio)

    obtenido = resumen.get("obtenido")
    total = resumen.get("total")
    nota = resumen.get("nota")
    buenas = resumen.get("buenas", 0)
    parciales = resumen.get("parciales", 0)
    incorrectas = resumen.get("incorrectas", 0)
    nivel = resumen.get("nivel", nivel_dificultad)

    st.divider()
    st.markdown('<div class="section-title">📊 Dashboard de resultados</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-sub">Resumen visual generado automáticamente por la IA.</div>', unsafe_allow_html=True)

    r1, r2, r3, r4, r5 = st.columns(5)

    with r1:
        metric_card("Punteo", f"{obtenido:g}/{total:g}" if obtenido is not None and total else "—", "Punteo detectado")

    with r2:
        metric_card("Nota", f"{nota:g}/100" if nota is not None else "—", "Escala final")

    with r3:
        metric_card("Buenas", str(buenas), "Respuestas correctas")

    with r4:
        metric_card("Parciales", str(parciales), "Respuestas incompletas")

    with r5:
        metric_card("Incorrectas", str(incorrectas), "Respuestas malas")

    g1, g2 = st.columns(2)

    with g1:
        st.markdown("### 🎯 Medidor de nota")
        st.plotly_chart(crear_gauge(nota), width="stretch")

    with g2:
        st.markdown("### 🧬 Distribución de respuestas")
        st.plotly_chart(crear_donut(buenas, parciales, incorrectas), width="stretch")

    st.markdown(
        f"""
        <div class="result-shell">
            <h2>📑 Informe de Calificación</h2>
            <p style="color:#94a3b8;">
                Nivel aplicado: {nivel}/10 · Motor: OCR/ICR + RAG + LLM · Imágenes: {len(st.session_state.imagenes_bytes)} · Marca: Grupo 8
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

    intro, preguntas = separar_preguntas(resultado_limpio)

    with st.expander("📌 Resumen general", expanded=True):
        st.markdown(intro)

    if preguntas:
        st.markdown("### 🧾 Detalle por pregunta")

        for i, bloque in enumerate(preguntas, start=1):
            icono_estado = estado_pregunta(bloque)
            titulo = f"Pregunta global {i}"

            match_titulo = re.search(r"###\s*(Pregunta global\s+\d+)", bloque, flags=re.IGNORECASE)
            if match_titulo:
                titulo = match_titulo.group(1)

            with st.expander(f"{icono_estado} {titulo}", expanded=False):
                st.markdown('<div class="question-card">', unsafe_allow_html=True)
                st.markdown(bloque)
                st.markdown('</div>', unsafe_allow_html=True)

    with st.expander("📄 Ver informe completo", expanded=False):
        st.markdown(resultado_limpio)

    # ========================================================
    # DATOS DEL REPORTE
    # ========================================================

    st.divider()
    st.markdown('<div class="section-title">🧾 Datos para reporte e historial</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-sub">La IA intenta llenar estos datos automáticamente. El docente puede verificarlos o corregirlos antes de guardar.</div>',
        unsafe_allow_html=True
    )

    paquete_actual = st.session_state.paquete_cache or {}
    metadatos_detectados = paquete_actual.get("metadatos") or {}

    if metadatos_detectados:
        with st.expander("🤖 Datos detectados automáticamente por OCR/ICR", expanded=False):
            st.json(metadatos_detectados)
    else:
        st.info("No se detectaron metadatos claros. El docente puede llenar los campos manualmente.")

    datos_base = st.session_state.datos_reporte or {}

    d1, d2, d3 = st.columns(3)

    with d1:
        estudiante_manual = st.text_input("Nombre del estudiante", value=datos_base.get("estudiante", ""))
        codigo_manual = st.text_input("Código / carné", value=datos_base.get("codigo", ""))

    with d2:
        curso_manual = st.text_input("Curso", value=datos_base.get("curso", ""))
        docente_manual = st.text_input("Docente", value=datos_base.get("docente", ""))

    with d3:
        titulo_manual = st.text_input("Título del examen", value=datos_base.get("titulo", ""))
        serie_manual = st.text_input("Serie del examen", value=datos_base.get("serie", ""))

    fecha_examen_manual = st.text_input(
        "Fecha del examen",
        value=datos_base.get("fecha_examen", ""),
        placeholder="Opcional"
    )

    st.session_state.datos_reporte = {
        "examen_id": st.session_state.ultimo_guardado_id,
        "estudiante": estudiante_manual,
        "codigo": codigo_manual,
        "curso": curso_manual,
        "docente": docente_manual,
        "titulo": titulo_manual,
        "serie": serie_manual,
        "fecha_examen": fecha_examen_manual,
    }

    datos_descarga = st.session_state.datos_reporte.copy()

    # ========================================================
    # DESCARGAS
    # ========================================================

    st.markdown("### 📥 Descargar reporte")

    try:
        pdf_bytes = generar_pdf_reporte(datos_descarga, resumen, resultado_limpio)
        word_bytes = generar_word_reporte(datos_descarga, resumen, resultado_limpio)
    except Exception as e:
        st.error("No se pudieron generar los archivos de descarga.")
        st.exception(e)
        pdf_bytes = None
        word_bytes = None

    nombre_pdf = nombre_archivo_reporte(
        "pdf",
        st.session_state.ultimo_guardado_id,
        estudiante_manual,
        curso_manual
    )

    nombre_word = nombre_archivo_reporte(
        "docx",
        st.session_state.ultimo_guardado_id,
        estudiante_manual,
        curso_manual
    )

    down1, down2 = st.columns(2)

    with down1:
        st.download_button(
            label="📄 Descargar PDF",
            data=pdf_bytes or b"",
            file_name=nombre_pdf,
            mime="application/pdf",
            disabled=pdf_bytes is None
        )

    with down2:
        st.download_button(
            label="📝 Descargar Word",
            data=word_bytes or b"",
            file_name=nombre_word,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            disabled=word_bytes is None
        )

    st.info(
        f"Formato de descarga: {nombre_archivo_reporte('pdf', st.session_state.ultimo_guardado_id, estudiante_manual, curso_manual)}"
    )

    # ========================================================
    # GUARDAR EN HISTORIAL
    # ========================================================

    st.divider()
    st.markdown('<div class="section-title">💾 Guardar en historial</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-sub">Guarda el resultado en PostgreSQL para consultarlo desde la página de historial.</div>',
        unsafe_allow_html=True
    )

    guardar = st.button("💾 Guardar calificación en historial")

    if guardar:
        try:
            paquete = st.session_state.paquete_cache or {}

            imagen_nombre, imagen_hash = resumen_archivos_examen(st.session_state.imagenes_bytes)

            conclusion_general = extraer_texto_seccion("Conclusión General", resultado_limpio)
            recomendaciones = extraer_texto_seccion("Recomendaciones de Estudio", resultado_limpio)

            examen_id = guardar_examen_calificado(
                estudiante_nombre=estudiante_manual or None,
                estudiante_codigo=codigo_manual or None,
                curso=curso_manual or None,
                docente=docente_manual or None,
                titulo_examen=titulo_manual or None,
                serie_examen=serie_manual or None,
                fecha_examen=fecha_examen_manual or None,

                archivo_nombre=imagen_nombre,
                archivo_tipo="imagenes" if len(st.session_state.imagenes_bytes) > 1 else "imagen",
                archivo_hash=imagen_hash,

                nivel_dificultad=nivel,

                nota_obtenida=obtenido,
                nota_maxima=total,
                nota_escala_100=nota,
                porcentaje=nota,

                preguntas_buenas=buenas,
                preguntas_parciales=parciales,
                preguntas_incorrectas=incorrectas,
                total_preguntas=buenas + parciales + incorrectas,

                conclusion_general=conclusion_general,
                fortalezas=None,
                debilidades=None,
                recomendaciones_estudio=recomendaciones,

                texto_examen_extraido=paquete.get("texto_examen"),
                metadatos_extraidos_json=paquete.get("metadatos"),
                rubrica_detectada_json=paquete.get("rubrica"),
                informe_markdown=resultado_limpio
            )

            st.session_state.ultimo_guardado_id = examen_id
            st.session_state.datos_reporte["examen_id"] = examen_id

            st.success(f"✅ Calificación guardada correctamente. ID del historial: {examen_id}")
            st.info("Ahora los próximos reportes descargados llevarán ese ID en el nombre del archivo.")

        except Exception as e:
            st.error("❌ No se pudo guardar en historial.")
            st.exception(e)

    if st.session_state.ultimo_guardado_id:
        nav1, nav2 = st.columns(2)

        with nav1:
            if st.button("📚 Ir al historial de calificaciones"):
                st.switch_page("pages/historial.py")

        with nav2:
            st.info(f"Último registro guardado: ID {st.session_state.ultimo_guardado_id}")

else:
    st.divider()
    st.markdown(
        """
        <div class="result-shell">
            <h2>🧠 Esperando examen</h2>
            <p style="color:#94a3b8;">
                Sube el material de referencia y una o varias imágenes del examen para iniciar la evaluación inteligente.
            </p>
            <span class="badge">1. Lectura OCR/ICR multi-imagen</span>
            <span class="badge">2. Búsqueda RAG</span>
            <span class="badge">3. Evaluación LLM</span>
            <span class="badge">4. Reporte PDF/Word</span>
            <span class="badge">5. Guardado en historial</span>
        </div>
        """,
        unsafe_allow_html=True
    )


st.markdown(
    """
    <div class="footer-note">
        Grupo 8 · EvaluaIA Neural · Sistema Inteligente Multimodal para la Calificación Autónoma de Exámenes
    </div>
    """,
    unsafe_allow_html=True
)