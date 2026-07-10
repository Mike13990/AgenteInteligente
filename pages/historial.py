import json
import re
from io import BytesIO
from datetime import datetime

import streamlit as st
import plotly.graph_objects as go

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from db import listar_historial, obtener_examen_por_id


# ============================================================
# CONFIGURACIÓN
# ============================================================

st.set_page_config(
    page_title="Historial · EvaluaIA Neural",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="collapsed",
)


# ============================================================
# CSS UI/UX
# ============================================================

st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;800;900&display=swap');

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
}

.stApp {
    background:
        radial-gradient(circle at 12% 10%, rgba(37,99,235,0.30), transparent 34%),
        radial-gradient(circle at 92% 5%, rgba(168,85,247,0.22), transparent 34%),
        radial-gradient(circle at 50% 100%, rgba(14,165,233,0.14), transparent 42%),
        linear-gradient(135deg, #020617 0%, #0f172a 52%, #111827 100%);
    color: #e5e7eb;
}

.block-container {
    padding-top: 1.1rem;
    padding-bottom: 2.6rem;
    max-width: 1550px;
}

hr {
    border-color: rgba(148,163,184,0.18);
}

.hero {
    position: relative;
    overflow: hidden;
    border-radius: 34px;
    padding: 38px;
    margin-bottom: 22px;
    border: 1px solid rgba(125,211,252,0.30);
    background:
        linear-gradient(135deg, rgba(15,23,42,0.96), rgba(30,41,59,0.80)),
        radial-gradient(circle at 18% 20%, rgba(59,130,246,0.34), transparent 35%),
        radial-gradient(circle at 88% 8%, rgba(168,85,247,0.28), transparent 38%);
    box-shadow: 0 30px 90px rgba(0,0,0,0.50);
}

.hero::before {
    content: "";
    position: absolute;
    inset: -3px;
    background: linear-gradient(90deg, transparent, rgba(125,211,252,0.16), transparent);
    transform: translateX(-110%);
    animation: shine 6s infinite;
}

@keyframes shine {
    0% { transform: translateX(-110%); }
    55% { transform: translateX(120%); }
    100% { transform: translateX(120%); }
}

.hero-title {
    position: relative;
    font-size: 54px;
    line-height: 1;
    font-weight: 900;
    letter-spacing: -1.6px;
    background: linear-gradient(90deg, #93c5fd, #c4b5fd, #67e8f9);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
}

.hero-sub {
    position: relative;
    margin-top: 14px;
    max-width: 980px;
    color: #cbd5e1;
    font-size: 17px;
    line-height: 1.55;
}

.badge {
    display: inline-flex;
    align-items: center;
    gap: 6px;
    padding: 8px 13px;
    margin: 14px 7px 0 0;
    border-radius: 999px;
    color: #dbeafe;
    background: rgba(37,99,235,0.15);
    border: 1px solid rgba(96,165,250,0.35);
    font-size: 13px;
    font-weight: 800;
}

.glass {
    border-radius: 26px;
    padding: 20px;
    border: 1px solid rgba(148,163,184,0.20);
    background: rgba(15,23,42,0.62);
    box-shadow: 0 18px 55px rgba(0,0,0,0.28);
}

.card {
    border-radius: 26px;
    padding: 22px;
    min-height: 138px;
    border: 1px solid rgba(148,163,184,0.20);
    background:
        linear-gradient(180deg, rgba(15,23,42,0.90), rgba(2,6,23,0.72));
    box-shadow: 0 18px 48px rgba(0,0,0,0.30);
}

.card:hover {
    transform: translateY(-2px);
    border-color: rgba(125,211,252,0.45);
    box-shadow: 0 24px 60px rgba(0,0,0,0.40);
    transition: 0.2s ease;
}

.card-label {
    color: #93c5fd;
    font-size: 13px;
    font-weight: 900;
    text-transform: uppercase;
    letter-spacing: .10em;
}

.card-value {
    margin-top: 8px;
    color: #f8fafc;
    font-size: 30px;
    font-weight: 900;
}

.card-desc {
    margin-top: 8px;
    color: #94a3b8;
    font-size: 14px;
    line-height: 1.45;
}

.section-title {
    font-size: 28px;
    font-weight: 900;
    color: #f8fafc;
    margin: 12px 0 6px 0;
}

.section-sub {
    color: #94a3b8;
    margin-bottom: 16px;
}

.record-card {
    border-radius: 26px;
    padding: 20px;
    margin-bottom: 16px;
    border: 1px solid rgba(148,163,184,0.20);
    background:
        linear-gradient(135deg, rgba(15,23,42,0.92), rgba(30,41,59,0.55)),
        radial-gradient(circle at 96% 0%, rgba(59,130,246,0.16), transparent 32%);
    box-shadow: 0 18px 48px rgba(0,0,0,0.32);
}

.record-title {
    font-size: 22px;
    font-weight: 900;
    color: #f8fafc;
    margin-bottom: 4px;
}

.record-meta {
    color: #94a3b8;
    font-size: 14px;
    line-height: 1.45;
}

.result-shell {
    border-radius: 28px;
    padding: 24px;
    border: 1px solid rgba(148,163,184,0.22);
    background:
        linear-gradient(180deg, rgba(15,23,42,0.88), rgba(2,6,23,0.74));
    box-shadow: 0 22px 72px rgba(0,0,0,0.38);
}

.status-pass {
    display:inline-block;
    padding:7px 11px;
    border-radius:999px;
    background:rgba(34,197,94,0.15);
    color:#bbf7d0;
    border:1px solid rgba(34,197,94,0.35);
    font-weight:900;
    font-size:12px;
}

.status-mid {
    display:inline-block;
    padding:7px 11px;
    border-radius:999px;
    background:rgba(234,179,8,0.15);
    color:#fef3c7;
    border:1px solid rgba(234,179,8,0.35);
    font-weight:900;
    font-size:12px;
}

.status-low {
    display:inline-block;
    padding:7px 11px;
    border-radius:999px;
    background:rgba(239,68,68,0.15);
    color:#fecaca;
    border:1px solid rgba(239,68,68,0.35);
    font-weight:900;
    font-size:12px;
}

.stButton > button {
    width: 100%;
    border-radius: 18px;
    padding: 14px 18px;
    font-weight: 900;
    color: white;
    border: 1px solid rgba(125,211,252,0.45);
    background: linear-gradient(90deg, #2563eb, #7c3aed);
    box-shadow: 0 15px 36px rgba(37,99,235,0.26);
}

.stButton > button:hover {
    transform: translateY(-1px);
    border-color: rgba(125,211,252,0.85);
    box-shadow: 0 20px 50px rgba(124,58,237,0.34);
}

.stDownloadButton > button {
    width: 100%;
    border-radius: 18px;
    padding: 14px 18px;
    font-weight: 900;
    color: white;
    background: linear-gradient(90deg, #0f766e, #2563eb);
    border: 1px solid rgba(125,211,252,0.35);
}

div[data-testid="stExpander"] {
    border: 1px solid rgba(148,163,184,0.20);
    border-radius: 18px;
    background: rgba(15,23,42,0.38);
}

.footer-note {
    color: #64748b;
    font-size: 13px;
    text-align: center;
    padding-top: 18px;
}
</style>
""",
    unsafe_allow_html=True,
)


# ============================================================
# SESSION STATE
# ============================================================

if "historial_id_abierto" not in st.session_state:
    st.session_state["historial_id_abierto"] = None


# ============================================================
# FUNCIONES
# ============================================================

def formato_num(valor):
    if valor is None:
        return "—"
    try:
        return f"{float(valor):g}"
    except Exception:
        return str(valor)


def formato_fecha(valor):
    if valor is None:
        return "—"
    try:
        if isinstance(valor, datetime):
            return valor.strftime("%Y-%m-%d %H:%M")
        return str(valor).replace("T", " ")[:19]
    except Exception:
        return str(valor)


def limpiar_para_archivo(texto: str) -> str:
    texto = texto or "SinDato"
    texto = texto.strip()
    texto = re.sub(r"[^\w\sáéíóúÁÉÍÓÚñÑ-]", "", texto)
    texto = re.sub(r"\s+", "_", texto)
    return texto[:60] or "SinDato"


def nombre_archivo_reporte(extension: str, examen: dict) -> str:
    fecha = datetime.now().strftime("%Y-%m-%d")
    examen_id = examen.get("id") or "SinID"
    estudiante = limpiar_para_archivo(examen.get("estudiante_nombre"))
    curso = limpiar_para_archivo(examen.get("curso"))
    return f"Reporte_ID{examen_id}_{estudiante}_{curso}_{fecha}.{extension}"


def limpiar_markdown(texto: str) -> str:
    return (
        str(texto or "")
        .replace("```markdown", "")
        .replace("```", "")
        .strip()
    )


def quitar_markdown_para_reporte(texto: str) -> str:
    texto = limpiar_markdown(texto)
    texto = re.sub(r"#{1,6}\s*", "", texto)
    texto = texto.replace("**", "")
    texto = texto.replace("*", "")
    texto = texto.replace("`", "")
    return texto.strip()


def parse_json_seguro(valor):
    if valor is None:
        return {}
    if isinstance(valor, (dict, list)):
        return valor
    try:
        return json.loads(valor)
    except Exception:
        return {"valor": str(valor)}


def nota_float(item):
    try:
        return float(item.get("nota_escala_100") or 0)
    except Exception:
        return 0.0


def estado_nota(nota):
    try:
        n = float(nota)
    except Exception:
        n = 0

    if n >= 80:
        return "Excelente", "status-pass"
    if n >= 60:
        return "Aceptable", "status-mid"
    return "Mejorar", "status-low"


def metric_card(label, value, desc):
    st.markdown(
        f"""
        <div class="card">
            <div class="card-label">{label}</div>
            <div class="card-value">{value}</div>
            <div class="card-desc">{desc}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def crear_donut(buenas, parciales, incorrectas):
    fig = go.Figure(
        data=[
            go.Pie(
                labels=["Buenas", "Parciales", "Incorrectas"],
                values=[buenas or 0, parciales or 0, incorrectas or 0],
                hole=0.62,
                textinfo="label+value",
                marker=dict(colors=["#22c55e", "#f59e0b", "#ef4444"]),
            )
        ]
    )
    fig.update_layout(
        height=310,
        margin=dict(l=10, r=10, t=20, b=10),
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        font=dict(color="#e5e7eb", size=14),
        showlegend=False,
    )
    return fig


def crear_gauge(nota):
    nota = nota or 0
    fig = go.Figure(
        go.Indicator(
            mode="gauge+number",
            value=float(nota),
            number={"suffix": "/100", "font": {"size": 34, "color": "#f8fafc"}},
            gauge={
                "axis": {"range": [0, 100], "tickcolor": "#94a3b8"},
                "bar": {"color": "#60a5fa"},
                "bgcolor": "rgba(15,23,42,0.55)",
                "borderwidth": 1,
                "bordercolor": "rgba(148,163,184,0.35)",
                "steps": [
                    {"range": [0, 60], "color": "rgba(239,68,68,0.25)"},
                    {"range": [60, 80], "color": "rgba(245,158,11,0.25)"},
                    {"range": [80, 100], "color": "rgba(34,197,94,0.25)"},
                ],
            },
        )
    )
    fig.update_layout(
        height=310,
        margin=dict(l=10, r=10, t=20, b=10),
        paper_bgcolor="rgba(0,0,0,0)",
        font=dict(color="#e5e7eb"),
    )
    return fig


def crear_barras_notas(items):
    ultimos = list(reversed(items[-10:]))
    etiquetas = [f"ID {x.get('id')}" for x in ultimos]
    valores = [nota_float(x) for x in ultimos]

    fig = go.Figure()
    fig.add_trace(
        go.Bar(
            x=etiquetas,
            y=valores,
            text=[f"{v:.1f}" for v in valores],
            textposition="outside",
        )
    )
    fig.update_layout(
        height=330,
        yaxis=dict(range=[0, 100], title="Nota /100"),
        xaxis=dict(title="Últimos registros"),
        margin=dict(l=10, r=10, t=25, b=10),
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(15,23,42,0.35)",
        font=dict(color="#e5e7eb"),
    )
    return fig


def marca_agua_pdf(canvas, doc):
    canvas.saveState()
    width, height = letter

    canvas.setFillColor(colors.Color(0.20, 0.36, 0.65, alpha=0.10))
    canvas.setFont("Helvetica-Bold", 46)
    canvas.translate(width / 2, height / 2)
    canvas.rotate(35)
    canvas.drawCentredString(0, 0.45 * inch, "GRUPO 8")
    canvas.setFont("Helvetica-Bold", 24)
    canvas.drawCentredString(0, 0, "EvaluaIA Neural")
    canvas.setFont("Helvetica", 16)
    canvas.drawCentredString(0, -0.35 * inch, "Sistema Inteligente de Calificación")
    canvas.restoreState()

    canvas.saveState()
    canvas.setFillColor(colors.HexColor("#64748B"))
    canvas.setFont("Helvetica", 8)
    canvas.drawString(0.7 * inch, 0.45 * inch, "Generado por Grupo 8 · EvaluaIA Neural")
    canvas.drawRightString(width - 0.7 * inch, 0.45 * inch, f"Página {doc.page}")
    canvas.restoreState()


def generar_pdf_reporte(examen: dict) -> bytes:
    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.7 * inch,
        leftMargin=0.7 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=22,
        leading=26,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#1E3A8A"),
        spaceAfter=12,
    )

    subtitle_style = ParagraphStyle(
        "CustomSubtitle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#475569"),
        spaceAfter=18,
    )

    h_style = ParagraphStyle(
        "CustomH",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#1E40AF"),
        spaceBefore=10,
        spaceAfter=8,
    )

    body_style = ParagraphStyle(
        "CustomBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#111827"),
        alignment=TA_LEFT,
    )

    story = []

    story.append(Paragraph("Reporte de Calificación Inteligente", title_style))
    story.append(Paragraph("GRUPO 8 · EvaluaIA Neural · Sistema Inteligente Multimodal", subtitle_style))

    info = [
        ["ID historial", str(examen.get("id") or "Sin ID")],
        ["Estudiante", examen.get("estudiante_nombre") or "No especificado"],
        ["Código / carné", examen.get("estudiante_codigo") or "No especificado"],
        ["Curso", examen.get("curso") or "No especificado"],
        ["Docente", examen.get("docente") or "No especificado"],
        ["Título del examen", examen.get("titulo_examen") or "No especificado"],
        ["Serie", examen.get("serie_examen") or "No especificado"],
        ["Fecha del examen", formato_fecha(examen.get("fecha_examen"))],
        ["Fecha de creación", formato_fecha(examen.get("creado_en"))],
        ["Fecha de descarga", datetime.now().strftime("%Y-%m-%d %H:%M")],
    ]

    tabla_info = Table(info, colWidths=[1.75 * inch, 4.85 * inch])
    tabla_info.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#DBEAFE")),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#1E3A8A")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("PADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )

    story.append(Paragraph("Datos generales", h_style))
    story.append(tabla_info)
    story.append(Spacer(1, 14))

    tabla_resumen = [
        ["Punteo", "Nota /100", "Nivel", "Buenas", "Parciales", "Incorrectas"],
        [
            f"{formato_num(examen.get('nota_obtenida'))}/{formato_num(examen.get('nota_maxima'))}",
            f"{formato_num(examen.get('nota_escala_100'))}/100",
            f"{formato_num(examen.get('nivel_dificultad'))}/10",
            str(examen.get("preguntas_buenas") or 0),
            str(examen.get("preguntas_parciales") or 0),
            str(examen.get("preguntas_incorrectas") or 0),
        ],
    ]

    t_resumen = Table(tabla_resumen, colWidths=[1.1 * inch] * 6)
    t_resumen.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor("#EFF6FF")),
                ("FONTNAME", (0, 1), (-1, 1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
                ("PADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )

    story.append(Paragraph("Resumen de calificación", h_style))
    story.append(t_resumen)
    story.append(Spacer(1, 14))

    story.append(Paragraph("Informe completo generado por IA", h_style))

    texto_limpio = quitar_markdown_para_reporte(examen.get("informe_markdown") or "No hay informe guardado.")

    for parrafo in texto_limpio.split("\n"):
        parrafo = parrafo.strip()
        if not parrafo:
            story.append(Spacer(1, 5))
        else:
            parrafo = parrafo.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(parrafo, body_style))

    doc.build(story, onFirstPage=marca_agua_pdf, onLaterPages=marca_agua_pdf)

    data = buffer.getvalue()
    buffer.close()
    return data


def generar_word_reporte(examen: dict) -> bytes:
    buffer = BytesIO()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    header = section.header
    h = header.paragraphs[0]
    h.text = "GRUPO 8 · EvaluaIA Neural · Sistema Inteligente de Calificación"
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer
    f = footer.paragraphs[0]
    f.text = "Generado por Grupo 8 · EvaluaIA Neural"
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER

    wm = doc.add_paragraph()
    wm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_wm = wm.add_run("GRUPO 8 · EvaluaIA Neural")
    run_wm.bold = True
    run_wm.font.size = Pt(28)
    run_wm.font.color.rgb = RGBColor(180, 190, 210)

    title = doc.add_heading("Reporte de Calificación Inteligente", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("GRUPO 8 · EvaluaIA Neural · Sistema Inteligente Multimodal")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Datos generales", level=1)

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    info = [
        ("ID historial", str(examen.get("id") or "Sin ID")),
        ("Estudiante", examen.get("estudiante_nombre") or "No especificado"),
        ("Código / carné", examen.get("estudiante_codigo") or "No especificado"),
        ("Curso", examen.get("curso") or "No especificado"),
        ("Docente", examen.get("docente") or "No especificado"),
        ("Título del examen", examen.get("titulo_examen") or "No especificado"),
        ("Serie", examen.get("serie_examen") or "No especificado"),
        ("Fecha del examen", formato_fecha(examen.get("fecha_examen"))),
        ("Fecha de creación", formato_fecha(examen.get("creado_en"))),
        ("Fecha de descarga", datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]

    for etiqueta, valor in info:
        row = table.add_row().cells
        row[0].text = etiqueta
        row[1].text = valor

    doc.add_heading("Resumen de calificación", level=1)

    table2 = doc.add_table(rows=2, cols=6)
    table2.style = "Table Grid"

    headers = ["Punteo", "Nota /100", "Nivel", "Buenas", "Parciales", "Incorrectas"]
    values = [
        f"{formato_num(examen.get('nota_obtenida'))}/{formato_num(examen.get('nota_maxima'))}",
        f"{formato_num(examen.get('nota_escala_100'))}/100",
        f"{formato_num(examen.get('nivel_dificultad'))}/10",
        str(examen.get("preguntas_buenas") or 0),
        str(examen.get("preguntas_parciales") or 0),
        str(examen.get("preguntas_incorrectas") or 0),
    ]

    for i, header_text in enumerate(headers):
        table2.rows[0].cells[i].text = header_text
        table2.rows[1].cells[i].text = values[i]

    doc.add_heading("Informe completo generado por IA", level=1)
    texto_limpio = quitar_markdown_para_reporte(examen.get("informe_markdown") or "No hay informe guardado.")

    for parrafo in texto_limpio.split("\n"):
        parrafo = parrafo.strip()
        if parrafo:
            doc.add_paragraph(parrafo)

    doc.save(buffer)
    data = buffer.getvalue()
    buffer.close()
    return data


# ============================================================
# ENCABEZADO
# ============================================================

st.markdown(
    """
    <div class="hero">
        <div class="hero-title">Historial Inteligente</div>
        <div class="hero-sub">
            Consulta, filtra, abre y descarga los reportes guardados por EvaluaIA Neural.
            Esta vista recupera las calificaciones desde PostgreSQL y permite exportar cada informe en PDF o Word.
        </div>
        <div>
            <span class="badge">🗄️ PostgreSQL</span>
            <span class="badge">📊 Dashboard</span>
            <span class="badge">📑 PDF / Word</span>
            <span class="badge">🏷️ Marca de agua Grupo 8</span>
            <span class="badge">🔎 Búsqueda avanzada</span>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)

top1, top2, top3 = st.columns([1.1, 1.1, 3.8])

with top1:
    if st.button("⬅️ Volver a calificar"):
        st.switch_page("app.py")

with top2:
    if st.button("🔄 Actualizar historial"):
        st.rerun()

with top3:
    st.markdown(
        """
        <div class="glass">
            <strong>Grupo 8 · EvaluaIA Neural</strong>
            <br>
            <span style="color:#94a3b8;">Los reportes descargados llevan ID, estudiante, curso, fecha y marca de agua.</span>
        </div>
        """,
        unsafe_allow_html=True,
    )


# ============================================================
# CARGA DE HISTORIAL
# ============================================================

try:
    historial = listar_historial()
except Exception as e:
    st.error(f"No se pudo cargar el historial desde la base de datos: {e}")
    st.stop()

if not historial:
    st.warning("Todavía no hay calificaciones guardadas. Primero califica un examen y presiona guardar en historial.")
    if st.button("🚀 Ir a calificar primer examen"):
        st.switch_page("app.py")
    st.stop()


# ============================================================
# MÉTRICAS GENERALES
# ============================================================

total_registros = len(historial)
notas = [nota_float(item) for item in historial if item.get("nota_escala_100") is not None]

promedio = sum(notas) / len(notas) if notas else 0
mejor = max(notas) if notas else 0
menor = min(notas) if notas else 0
aprobados = len([n for n in notas if n >= 60])

m1, m2, m3, m4, m5 = st.columns(5)

with m1:
    metric_card("Registros", str(total_registros), "Exámenes guardados")

with m2:
    metric_card("Promedio", f"{promedio:.1f}/100", "Promedio general")

with m3:
    metric_card("Mejor nota", f"{mejor:.1f}/100", "Nota más alta")

with m4:
    metric_card("Nota menor", f"{menor:.1f}/100", "Nota más baja")

with m5:
    metric_card("Aprobados", str(aprobados), "Notas iguales o mayores a 60")


# ============================================================
# GRÁFICA GENERAL
# ============================================================

st.divider()
st.markdown('<div class="section-title">📈 Vista general</div>', unsafe_allow_html=True)
st.markdown('<div class="section-sub">Últimos registros guardados y su nota en escala de 100.</div>', unsafe_allow_html=True)

st.plotly_chart(crear_barras_notas(historial), width="stretch")


# ============================================================
# FILTROS
# ============================================================

st.divider()
st.markdown('<div class="section-title">🔎 Búsqueda avanzada</div>', unsafe_allow_html=True)
st.markdown('<div class="section-sub">Filtra por estudiante, curso, ID, rango de nota o nivel de dificultad.</div>', unsafe_allow_html=True)

f1, f2, f3, f4, f5 = st.columns([1.25, 1.25, 0.7, 1, 1])

with f1:
    filtro_estudiante = st.text_input("Buscar estudiante", "")

with f2:
    filtro_curso = st.text_input("Buscar curso", "")

with f3:
    filtro_id = st.text_input("ID", "")

with f4:
    nota_minima = st.slider("Nota mínima", 0, 100, 0)

with f5:
    ordenar = st.selectbox(
        "Ordenar por",
        [
            "Más recientes",
            "Mejor nota",
            "Menor nota",
            "Estudiante A-Z",
            "Curso A-Z",
        ],
    )


filtrados = []

for item in historial:
    estudiante = str(item.get("estudiante_nombre") or "").lower()
    curso = str(item.get("curso") or "").lower()
    item_id = str(item.get("id") or "")
    nota = nota_float(item)

    cumple_estudiante = filtro_estudiante.lower() in estudiante if filtro_estudiante else True
    cumple_curso = filtro_curso.lower() in curso if filtro_curso else True
    cumple_id = filtro_id in item_id if filtro_id else True
    cumple_nota = nota >= nota_minima

    if cumple_estudiante and cumple_curso and cumple_id and cumple_nota:
        filtrados.append(item)


if ordenar == "Mejor nota":
    filtrados.sort(key=lambda x: nota_float(x), reverse=True)
elif ordenar == "Menor nota":
    filtrados.sort(key=lambda x: nota_float(x))
elif ordenar == "Estudiante A-Z":
    filtrados.sort(key=lambda x: str(x.get("estudiante_nombre") or "").lower())
elif ordenar == "Curso A-Z":
    filtrados.sort(key=lambda x: str(x.get("curso") or "").lower())
else:
    filtrados.sort(key=lambda x: str(x.get("creado_en") or ""), reverse=True)


st.caption(f"Mostrando {len(filtrados)} de {len(historial)} registros.")


# ============================================================
# LISTA DE REGISTROS
# ============================================================

st.divider()
st.markdown('<div class="section-title">📚 Registros guardados</div>', unsafe_allow_html=True)

if not filtrados:
    st.info("No hay registros que coincidan con los filtros actuales.")
    st.stop()


for item in filtrados:
    item_id = item.get("id")
    estudiante = item.get("estudiante_nombre") or "Sin nombre"
    curso = item.get("curso") or "Sin curso"
    titulo = item.get("titulo_examen") or "Examen sin título"
    nota = formato_num(item.get("nota_escala_100"))
    nota_valor = nota_float(item)
    obtenido = formato_num(item.get("nota_obtenida"))
    maxima = formato_num(item.get("nota_maxima"))
    nivel = formato_num(item.get("nivel_dificultad"))
    fecha = formato_fecha(item.get("creado_en"))
    estado_texto, estado_clase = estado_nota(nota_valor)

    st.markdown(
        f"""
        <div class="record-card">
            <div class="record-title">🧾 ID {item_id} · {estudiante}</div>
            <div class="record-meta">
                <strong>Curso:</strong> {curso} ·
                <strong>Examen:</strong> {titulo} ·
                <strong>Nota:</strong> {nota}/100 ·
                <strong>Punteo:</strong> {obtenido}/{maxima} ·
                <strong>Nivel:</strong> {nivel}/10 ·
                <strong>Fecha:</strong> {fecha}
            </div>
            <br>
            <span class="{estado_clase}">{estado_texto}</span>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.expander(f"📖 Abrir detalles del registro ID {item_id}", expanded=False):
        c1, c2, c3, c4 = st.columns(4)

        with c1:
            st.write(f"**Estudiante:** {estudiante}")
            st.write(f"**Código:** {item.get('estudiante_codigo') or 'No registrado'}")
            st.write(f"**Curso:** {curso}")

        with c2:
            st.write(f"**Docente:** {item.get('docente') or 'No registrado'}")
            st.write(f"**Examen:** {titulo}")
            st.write(f"**Serie:** {item.get('serie_examen') or 'No registrada'}")

        with c3:
            st.write(f"**Punteo:** {obtenido}/{maxima}")
            st.write(f"**Nota:** {nota}/100")
            st.write(f"**Nivel:** {nivel}/10")

        with c4:
            st.write(f"**Buenas:** {item.get('preguntas_buenas') or 0}")
            st.write(f"**Parciales:** {item.get('preguntas_parciales') or 0}")
            st.write(f"**Incorrectas:** {item.get('preguntas_incorrectas') or 0}")
            st.write(f"**Total:** {item.get('total_preguntas') or 0}")

        g1, g2 = st.columns(2)

        with g1:
            st.markdown("### 🎯 Nota")
            st.plotly_chart(crear_gauge(nota_valor), width="stretch")

        with g2:
            st.markdown("### 🧬 Respuestas")
            st.plotly_chart(
                crear_donut(
                    item.get("preguntas_buenas") or 0,
                    item.get("preguntas_parciales") or 0,
                    item.get("preguntas_incorrectas") or 0,
                ),
                width="stretch",
            )

        b1, b2 = st.columns([1, 3])

        with b1:
            if st.button(f"📑 Cargar informe ID {item_id}", key=f"abrir_{item_id}"):
                st.session_state["historial_id_abierto"] = item_id

        with b2:
            st.info("Carga el informe completo para ver el texto extraído, rúbrica y descargas PDF/Word.")

        if st.session_state.get("historial_id_abierto") == item_id:
            try:
                examen = obtener_examen_por_id(item_id)

                if not examen:
                    st.error("No se encontró el examen seleccionado.")
                    continue

                informe = limpiar_markdown(examen.get("informe_markdown") or "No hay informe guardado.")

                st.markdown("### 📑 Informe completo")
                st.markdown('<div class="result-shell">', unsafe_allow_html=True)
                st.markdown(informe)
                st.markdown("</div>", unsafe_allow_html=True)

                st.markdown("### 📥 Descargar este reporte")

                pdf_bytes = generar_pdf_reporte(examen)
                word_bytes = generar_word_reporte(examen)

                d1, d2 = st.columns(2)

                with d1:
                    st.download_button(
                        label="📄 Descargar PDF",
                        data=pdf_bytes,
                        file_name=nombre_archivo_reporte("pdf", examen),
                        mime="application/pdf",
                        key=f"pdf_{item_id}",
                    )

                with d2:
                    st.download_button(
                        label="📝 Descargar Word",
                        data=word_bytes,
                        file_name=nombre_archivo_reporte("docx", examen),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"word_{item_id}",
                    )


                st.info(f"Formato: {nombre_archivo_reporte('pdf', examen)}")

                with st.expander("👁️ Texto extraído del examen", expanded=False):
                    st.text(examen.get("texto_examen_extraido") or "No hay texto extraído guardado.")

                with st.expander("🧾 Rúbrica detectada", expanded=False):
                    st.json(parse_json_seguro(examen.get("rubrica_detectada_json")))

            except Exception as e:
                st.error(f"No se pudo abrir el informe: {e}")


st.markdown(
    """
    <div class="footer-note">
        Grupo 8 · EvaluaIA Neural · Historial de calificaciones con reportes PDF y Word
    </div>
    """,
    unsafe_allow_html=True,
)