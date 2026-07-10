import os
import re
import json
import hashlib
import time
from pathlib import Path
from dataclasses import dataclass, asdict
from typing import Optional, Union, List, Dict, Any

import chromadb
from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction
from dotenv import load_dotenv
from google import genai
from google.genai import types
from pypdf import PdfReader


# ============================================================
# CONFIGURACIÓN
# ============================================================

load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")

if not GEMINI_API_KEY:
    raise ValueError("Falta GEMINI_API_KEY o GOOGLE_API_KEY en el archivo .env")

os.environ["GEMINI_API_KEY"] = GEMINI_API_KEY
os.environ["GOOGLE_API_KEY"] = GEMINI_API_KEY

MODELO_PRINCIPAL = os.getenv("MODELO_PRINCIPAL", "gemini-2.5-flash")
MODELO_FALLBACK = os.getenv("MODELO_FALLBACK", "gemini-2.5-flash-lite")


# ============================================================
# CARPETAS CACHE
# ============================================================

CACHE_DIR = Path(".cache_examenes")
CACHE_DIR.mkdir(parents=True, exist_ok=True)

RAG_DIR = Path(".rag_cache")
RAG_DIR.mkdir(parents=True, exist_ok=True)


# ============================================================
# DATOS
# ============================================================

@dataclass
class PreguntaExamen:
    id_global: int
    serie: str
    numero_local: Optional[int]
    bloque: str
    valor: float


# ============================================================
# UTILIDADES GENERALES
# ============================================================

def normalizar_rutas(rutas: Union[str, Path, List[Union[str, Path]]]) -> List[str]:
    """
    Permite que app.py mande una sola imagen o varias imágenes.
    Si recibe string/path, lo convierte a lista.
    Si recibe lista, limpia valores vacíos.
    """
    if rutas is None:
        return []

    if isinstance(rutas, (str, Path)):
        return [str(rutas)]

    return [str(r) for r in rutas if r]


def hash_archivo(ruta: Union[str, Path]) -> str:
    h = hashlib.sha256()
    with open(ruta, "rb") as f:
        for bloque in iter(lambda: f.read(1024 * 1024), b""):
            h.update(bloque)
    return h.hexdigest()


def hash_archivos(rutas: List[Union[str, Path]]) -> str:
    h = hashlib.sha256()
    for ruta in rutas:
        ruta = str(ruta)
        h.update(Path(ruta).name.encode("utf-8", errors="ignore"))
        with open(ruta, "rb") as f:
            for bloque in iter(lambda: f.read(1024 * 1024), b""):
                h.update(bloque)
    return h.hexdigest()


def cargar_json_cache(nombre: str):
    ruta = CACHE_DIR / nombre
    if not ruta.exists():
        return None

    try:
        with open(ruta, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        print(f"No se pudo leer cache {nombre}: {e}")
        return None


def guardar_json_cache(nombre: str, data: dict):
    ruta = CACHE_DIR / nombre
    try:
        with open(ruta, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"No se pudo guardar cache {nombre}: {e}")


def limpiar_texto(texto: Any) -> str:
    if texto is None:
        return ""
    return str(texto).strip()


def normalizar_numero(valor):
    if valor is None:
        return None

    try:
        return float(str(valor).strip().replace(",", "."))
    except Exception:
        return None


def buscar_numero_patron(texto: str, patrones: list):
    for patron in patrones:
        match = re.search(patron, texto, flags=re.IGNORECASE | re.DOTALL)
        if match:
            numero = normalizar_numero(match.group(1))
            if numero is not None:
                return numero
    return None


# ============================================================
# GEMINI DIRECTO CON REINTENTOS
# ============================================================

def llamar_gemini(prompt: str, imagen_bytes: bytes = None, mime_type: str = None) -> str:
    cliente = genai.Client(api_key=GEMINI_API_KEY)

    modelos = [MODELO_PRINCIPAL, MODELO_FALLBACK]

    for modelo in modelos:
        for intento in range(3):
            try:
                if imagen_bytes and mime_type:
                    contents = [
                        types.Part.from_bytes(
                            data=imagen_bytes,
                            mime_type=mime_type
                        ),
                        prompt,
                    ]
                else:
                    contents = [prompt]

                respuesta = cliente.models.generate_content(
                    model=modelo,
                    contents=contents,
                )

                texto = respuesta.text.strip() if respuesta and respuesta.text else ""
                if texto:
                    return texto

            except Exception as e:
                print(f"[Gemini] {modelo} intento {intento + 1}/3 falló: {e}")
                time.sleep(1.5)

    return ""


# ============================================================
# OCR / ICR
# ============================================================

OCR_PROMPT = """
Eres un sistema experto en OCR e ICR para exámenes académicos.

Debes leer TODA la imagen:
- encabezado
- nombre del estudiante
- código, carné o clave del estudiante
- curso
- docente
- grado, sección o carrera si aparece
- fecha del examen
- título o tipo de examen
- serie del examen
- instrucciones
- tabla de puntuación
- series
- preguntas
- respuestas del estudiante
- valores de cada serie
- valores de cada pregunta

REGLAS IMPORTANTES:
1. No inventes preguntas.
2. No inventes respuestas.
3. No inventes datos del estudiante.
4. Si un dato no aparece o no se entiende, escribe "No detectado".
5. Si hay tabla de puntuación, úsala como fuente principal.
6. Si arriba dice algo como 15/100 pero la tabla indica Serie 1, Serie 2 y Total, usa la tabla.
7. Conserva el orden real del examen.
8. Si algo no se entiende, escribe "No legible".
9. Extrae preguntas y respuestas de forma clara.
10. Si una pregunta dice "(1 punto)", "(2 puntos)", etc., extrae ese valor.
11. No confundas escala de calificación con valor real del examen.
12. No ignores textos pequeños donde diga "Valor 10 puntos", "Valor 5 puntos", "1 punto", etc.
13. Si esta imagen es una página intermedia de un examen de varias páginas, extrae solo lo que aparezca en esta imagen.
14. Mantén el texto fiel a la imagen.

OBLIGATORIO PARA METADATOS:
Devuelve esta sección incluso si no detectas datos:

METADATOS_DETECTADOS:
Estudiante: ... / No detectado
Código o carné: ... / No detectado
Curso: ... / No detectado
Docente: ... / No detectado
Título del examen: ... / No detectado
Serie del examen: ... / No detectado
Fecha del examen: ... / No detectado
Grado o sección: ... / No detectado

OBLIGATORIO PARA PUNTEO:
- Si ves "PRIMERA SERIE: Valor 10 puntos", escribe exactamente:
Serie 1: 10
- Si ves "SEGUNDA SERIE: Valor 5 puntos", escribe exactamente:
Serie 2: 5
- Si ves "Total 15", escribe exactamente:
Total: 15
- Si una pregunta dice "(1 punto)", escribe:
Puntos indicados en la pregunta: 1
- Nunca ignores los valores aunque estén pequeños o al lado del título.

Devuelve exactamente este formato:

METADATOS_DETECTADOS:
Estudiante: ... / No detectado
Código o carné: ... / No detectado
Curso: ... / No detectado
Docente: ... / No detectado
Título del examen: ... / No detectado
Serie del examen: ... / No detectado
Fecha del examen: ... / No detectado
Grado o sección: ... / No detectado

RUBRICA_DETECTADA:
Serie 1: ... / No especificado
Serie 2: ... / No especificado
Total: ... / No especificado
Fuente de puntuación: tabla / instrucciones / encabezado / no especificado

PREGUNTAS_EXTRAIDAS:

Serie: Primera serie
Pregunta 1:
Pregunta: ...
Respuesta del estudiante: ...
Puntos indicados en la pregunta: ... / No especificado

Serie: Segunda serie
Pregunta 1:
Pregunta: ...
Respuesta del estudiante: ...
Puntos indicados en la pregunta: ... / No especificado

Si no puedes leer el examen, responde:
NO HAY EXAMEN LEGIBLE.
"""


def obtener_mime_imagen(ruta_imagen: str) -> str:
    extension = Path(ruta_imagen).suffix.lower()

    if extension == ".png":
        return "image/png"

    if extension in [".jpg", ".jpeg"]:
        return "image/jpeg"

    if extension == ".webp":
        return "image/webp"

    return "image/jpeg"


def extraer_texto_imagen_gemini(ruta_imagen: str) -> str:
    imagen_hash = hash_archivo(ruta_imagen)
    cache_name = f"ocr_{imagen_hash}.json"

    cache = cargar_json_cache(cache_name)
    if cache and cache.get("texto"):
        print(f"✅ OCR desde cache: {Path(ruta_imagen).name}")
        return cache["texto"]

    mime_type = obtener_mime_imagen(ruta_imagen)

    try:
        with open(ruta_imagen, "rb") as f:
            imagen_bytes = f.read()
    except Exception as e:
        print(f"No se pudo abrir imagen {ruta_imagen}: {e}")
        return ""

    texto = llamar_gemini(
        OCR_PROMPT,
        imagen_bytes=imagen_bytes,
        mime_type=mime_type
    )

    if texto:
        guardar_json_cache(cache_name, {"texto": texto})

    return texto


def extraer_texto_imagen(ruta_imagen: str) -> str:
    return extraer_texto_imagen_gemini(ruta_imagen)


def extraer_texto_varias_imagenes(rutas_imagenes: List[str]) -> str:
    """
    Une varias imágenes del mismo examen como si fueran varias páginas.
    """
    secciones = []

    for i, ruta in enumerate(rutas_imagenes, start=1):
        texto_pagina = extraer_texto_imagen_gemini(ruta)

        if not texto_pagina.strip():
            return f"ERROR_IMAGEN_{i}: No se pudo extraer texto de la imagen {Path(ruta).name}."

        if "NO HAY EXAMEN LEGIBLE" in texto_pagina.upper():
            return f"ERROR_IMAGEN_{i}: La imagen {Path(ruta).name} no contiene un examen legible."

        secciones.append(f"""
========================
PÁGINA / IMAGEN {i}
ARCHIVO: {Path(ruta).name}
========================
{texto_pagina}
""")

    return "\n\n".join(secciones).strip()


# ============================================================
# METADATOS DEL EXAMEN
# ============================================================

def limpiar_valor_metadato(valor: str) -> str:
    valor = limpiar_texto(valor)
    valor = re.sub(r"\s+", " ", valor).strip()
    valor = valor.strip("-:;,. ")

    if not valor:
        return ""

    bajos = valor.lower()
    invalidos = [
        "no detectado",
        "no especificado",
        "no legible",
        "no aparece",
        "n/a",
        "na",
        "none",
        "null",
    ]

    if bajos in invalidos:
        return ""

    return valor


def extraer_linea_metadato(texto: str, etiquetas: List[str]) -> str:
    for etiqueta in etiquetas:
        patron = rf"{re.escape(etiqueta)}\s*:\s*(.+)"
        match = re.search(patron, texto, flags=re.IGNORECASE)
        if match:
            linea = match.group(1).splitlines()[0]
            return limpiar_valor_metadato(linea)
    return ""


def extraer_metadatos_examen(texto_examen: str) -> Dict[str, str]:
    """
    Extrae datos para rellenar automáticamente los campos del reporte.
    El docente puede verificarlos o editarlos desde app.py.
    """
    texto = texto_examen or ""

    metadatos = {
        "estudiante": extraer_linea_metadato(texto, ["Estudiante", "Nombre del estudiante", "Nombre"]),
        "codigo": extraer_linea_metadato(texto, ["Código o carné", "Codigo o carne", "Código", "Codigo", "Carné", "Carne", "Clave"]),
        "curso": extraer_linea_metadato(texto, ["Curso", "Materia", "Asignatura"]),
        "docente": extraer_linea_metadato(texto, ["Docente", "Profesor", "Catedrático", "Catedratico", "Maestro"]),
        "titulo": extraer_linea_metadato(texto, ["Título del examen", "Titulo del examen", "Título", "Titulo", "Examen"]),
        "serie": extraer_linea_metadato(texto, ["Serie del examen", "Serie", "Versión", "Version"]),
        "fecha_examen": extraer_linea_metadato(texto, ["Fecha del examen", "Fecha"]),
        "grado_seccion": extraer_linea_metadato(texto, ["Grado o sección", "Grado o seccion", "Grado", "Sección", "Seccion"]),
    }

    # Fallbacks por patrones frecuentes de encabezado.
    if not metadatos["estudiante"]:
        match = re.search(r"(?:Alumno|Alumna|Estudiante|Nombre)\s*[:\-]\s*([^\n\r]+)", texto, flags=re.IGNORECASE)
        if match:
            metadatos["estudiante"] = limpiar_valor_metadato(match.group(1))

    if not metadatos["fecha_examen"]:
        match = re.search(r"\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b", texto)
        if match:
            metadatos["fecha_examen"] = match.group(1)

    return metadatos


# ============================================================
# LECTURA DE DOCUMENTOS DE REFERENCIA
# ============================================================

def leer_archivo_contexto(ruta: str) -> str:
    extension = Path(ruta).suffix.lower()

    try:
        if extension == ".pdf":
            reader = PdfReader(ruta)
            texto = ""

            for page_num, page in enumerate(reader.pages, start=1):
                contenido = page.extract_text() or ""
                if contenido.strip():
                    texto += f"\n\n[Página {page_num}]\n{contenido}"

            return texto.strip()

        if extension in [".txt", ".md"]:
            with open(ruta, "r", encoding="utf-8", errors="ignore") as f:
                return f.read().strip()

        if extension == ".docx":
            from docx import Document
            doc = Document(ruta)
            partes = []

            for p in doc.paragraphs:
                if p.text.strip():
                    partes.append(p.text.strip())

            for tabla in doc.tables:
                for fila in tabla.rows:
                    celdas = [c.text.strip() for c in fila.cells if c.text.strip()]
                    if celdas:
                        partes.append(" | ".join(celdas))

            return "\n".join(partes).strip()

        print(f"Formato de contexto no soportado: {extension}")
        return ""

    except Exception as e:
        print(f"Error leyendo contexto {ruta}: {e}")
        return ""


def leer_contextos(rutas_contexto: list) -> str:
    rutas_contexto = normalizar_rutas(rutas_contexto)

    if not rutas_contexto:
        return ""

    contexto_hash = hash_archivos(rutas_contexto)
    cache_name = f"contexto_{contexto_hash}.json"

    cache = cargar_json_cache(cache_name)
    if cache and cache.get("texto"):
        print("✅ Contexto desde cache.")
        return cache["texto"]

    textos = []

    for i, ruta in enumerate(rutas_contexto, start=1):
        texto = leer_archivo_contexto(ruta)

        if texto.strip():
            textos.append(
                f"""
========================
DOCUMENTO {i}
ARCHIVO: {Path(ruta).name}
========================
{texto}
"""
            )
        else:
            print(f"Documento sin texto útil: {Path(ruta).name}")

    texto_final = "\n".join(textos).strip()

    if texto_final:
        guardar_json_cache(cache_name, {"texto": texto_final})

    return texto_final


# ============================================================
# RÚBRICA Y PUNTEO
# ============================================================

def contar_preguntas_ocr(texto_examen: str) -> int:
    # Cuenta bloques tipo Pregunta 1:, Pregunta 2:, etc.
    return len(re.findall(r"Pregunta\s+\d+\s*:", texto_examen, flags=re.IGNORECASE))


def detectar_rubrica(texto_examen: str) -> dict:
    texto = texto_examen.replace("\n", " ")

    serie1 = buscar_numero_patron(texto, [
        r"Serie\s*1\s*:\s*([0-9]+(?:[.,][0-9]+)?)",
        r"Primera\s+serie\s*:\s*([0-9]+(?:[.,][0-9]+)?)",
        r"Primera\s+serie.*?Valor\s*[:=]?\s*([0-9]+(?:[.,][0-9]+)?)",
        r"PRIMERA\s+SERIE\s*:\s*Valor\s*([0-9]+(?:[.,][0-9]+)?)",
    ])

    serie2 = buscar_numero_patron(texto, [
        r"Serie\s*2\s*:\s*([0-9]+(?:[.,][0-9]+)?)",
        r"Segunda\s+serie\s*:\s*([0-9]+(?:[.,][0-9]+)?)",
        r"Segunda\s+serie.*?Valor\s*[:=]?\s*([0-9]+(?:[.,][0-9]+)?)",
        r"SEGUNDA\s+SERIE\s*:\s*Valor\s*([0-9]+(?:[.,][0-9]+)?)",
    ])

    total = buscar_numero_patron(texto, [
        r"Total\s*:\s*([0-9]+(?:[.,][0-9]+)?)",
        r"Total\s+([0-9]+(?:[.,][0-9]+)?)",
        r"valor\s+total\s*[:=]?\s*([0-9]+(?:[.,][0-9]+)?)",
        r"Punteo\s+total\s*[:=]?\s*([0-9]+(?:[.,][0-9]+)?)",
    ])

    if serie1 is not None or serie2 is not None:
        serie1 = serie1 if serie1 is not None else 0
        serie2 = serie2 if serie2 is not None else 0
        total_series = serie1 + serie2

        if total is None or abs(total - total_series) <= max(1, total_series * 0.10):
            total = total_series

        return {
            "serie1": serie1,
            "serie2": serie2,
            "total": total or total_series,
            "fuente": "rúbrica detectada en examen"
        }

    if total is not None and 0 < total <= 100:
        return {
            "serie1": None,
            "serie2": None,
            "total": total,
            "fuente": "total general detectado"
        }

    cantidad_preguntas = contar_preguntas_ocr(texto_examen)

    if cantidad_preguntas > 0:
        return {
            "serie1": None,
            "serie2": None,
            "total": float(cantidad_preguntas),
            "fuente": f"no se detectó tabla de puntuación; se asumió 1 punto por cada una de las {cantidad_preguntas} preguntas"
        }

    return {
        "serie1": None,
        "serie2": None,
        "total": 100.0,
        "fuente": "no se detectó puntuación; se asignó escala de 100"
    }


# ============================================================
# PREGUNTAS
# ============================================================

def detectar_serie(linea: str):
    l = linea.lower()

    if "primera serie" in l or "serie: primera" in l or "serie 1" in l:
        return "Serie 1"

    if "segunda serie" in l or "serie: segunda" in l or "serie 2" in l:
        return "Serie 2"

    if "tercera serie" in l or "serie: tercera" in l or "serie 3" in l:
        return "Serie 3"

    if "cuarta serie" in l or "serie: cuarta" in l or "serie 4" in l:
        return "Serie 4"

    return None


def dividir_preguntas_con_series(texto_examen: str) -> list:
    lineas = texto_examen.splitlines()
    preguntas = []
    serie_actual = "Sin serie"
    bloque_actual = []
    numero_actual = None

    for linea in lineas:
        serie_detectada = detectar_serie(linea)

        if serie_detectada:
            if bloque_actual:
                preguntas.append({
                    "serie": serie_actual,
                    "numero_local": numero_actual,
                    "bloque": "\n".join(bloque_actual).strip()
                })
                bloque_actual = []
                numero_actual = None

            serie_actual = serie_detectada
            continue

        match_pregunta = re.match(
            r"\s*Pregunta\s+(\d+)\s*:",
            linea,
            flags=re.IGNORECASE
        )

        if match_pregunta:
            if bloque_actual:
                preguntas.append({
                    "serie": serie_actual,
                    "numero_local": numero_actual,
                    "bloque": "\n".join(bloque_actual).strip()
                })

            numero_actual = int(match_pregunta.group(1))
            bloque_actual = [linea]
        else:
            if bloque_actual:
                bloque_actual.append(linea)

    if bloque_actual:
        preguntas.append({
            "serie": serie_actual,
            "numero_local": numero_actual,
            "bloque": "\n".join(bloque_actual).strip()
        })

    preguntas_limpias = [
        p for p in preguntas
        if "Pregunta:" in p["bloque"] and "Respuesta del estudiante:" in p["bloque"]
    ]

    if not preguntas_limpias and texto_examen.strip():
        return [{
            "serie": "Sin serie",
            "numero_local": 1,
            "bloque": texto_examen
        }]

    return preguntas_limpias


def extraer_valor_pregunta(bloque: str):
    patrones = [
        r"Puntos\s+indicados\s+en\s+la\s+pregunta\s*:\s*([0-9]+(?:[.,][0-9]+)?)",
        r"\(\s*([0-9]+(?:[.,][0-9]+)?)\s*punto[s]?\s*\)",
        r"Valor\s*:\s*([0-9]+(?:[.,][0-9]+)?)\s*punto[s]?",
        r"Vale\s*([0-9]+(?:[.,][0-9]+)?)\s*punto[s]?",
    ]

    for patron in patrones:
        match = re.search(patron, bloque, flags=re.IGNORECASE)
        if match:
            valor = normalizar_numero(match.group(1))
            if valor is not None and valor > 0:
                return valor

    return None


def asignar_puntajes(preguntas_raw: list, rubrica: dict) -> list:
    total = float(rubrica["total"])
    cantidad = max(len(preguntas_raw), 1)

    series = {}
    for p in preguntas_raw:
        series[p["serie"]] = series.get(p["serie"], 0) + 1

    valores_detectados = []
    for p in preguntas_raw:
        valor_individual = extraer_valor_pregunta(p["bloque"])
        valores_detectados.append(valor_individual)

    suma_valores_individuales = sum(v or 0 for v in valores_detectados)

    usar_valores_individuales = (
        any(v is not None for v in valores_detectados)
        and abs(suma_valores_individuales - total) <= max(1, total * 0.10)
    )

    preguntas = []

    for idx, p in enumerate(preguntas_raw, start=1):
        serie = p["serie"]
        valor_individual = valores_detectados[idx - 1]

        if usar_valores_individuales and valor_individual is not None:
            valor = valor_individual

        elif (rubrica.get("serie1") is not None or rubrica.get("serie2") is not None) and serie in ["Serie 1", "Serie 2"]:
            if serie == "Serie 1" and series.get("Serie 1", 0) > 0:
                valor = float(rubrica.get("serie1") or 0) / series["Serie 1"]
            elif serie == "Serie 2" and series.get("Serie 2", 0) > 0:
                valor = float(rubrica.get("serie2") or 0) / series["Serie 2"]
            else:
                valor = total / cantidad

        else:
            valor = total / cantidad

        preguntas.append(PreguntaExamen(
            id_global=idx,
            serie=serie,
            numero_local=p.get("numero_local"),
            bloque=p.get("bloque", ""),
            valor=round(float(valor), 2)
        ))

    return preguntas


def construir_tabla_rubrica(preguntas: list, rubrica: dict) -> str:
    filas = [
        f"Pregunta global {p.id_global} | {p.serie} | Pregunta {p.numero_local}: valor máximo {p.valor} puntos"
        for p in preguntas
    ]

    suma = round(sum(p.valor for p in preguntas), 2)

    return f"""
Fuente de puntuación: {rubrica["fuente"]}
Serie 1: {rubrica["serie1"] if rubrica["serie1"] is not None else "No especificado"}
Serie 2: {rubrica["serie2"] if rubrica["serie2"] is not None else "No especificado"}
Total del examen detectado: {rubrica["total"]}
Suma real de valores por pregunta: {suma}

Distribución por pregunta:
{chr(10).join(filas)}
"""


# ============================================================
# RAG
# ============================================================

def dividir_en_chunks(texto: str, tamano_chunk=900, solapamiento=120) -> list:
    texto = " ".join((texto or "").split())

    if not texto:
        return []

    chunks = []
    inicio = 0
    paso = max(tamano_chunk - solapamiento, 100)

    while inicio < len(texto):
        chunk = texto[inicio:inicio + tamano_chunk].strip()

        if chunk:
            chunks.append(chunk)

        inicio += paso

    return chunks


def construir_rag_cache(rutas_contexto: list, texto_contexto: str):
    rutas_contexto = normalizar_rutas(rutas_contexto)

    if not rutas_contexto or not texto_contexto.strip():
        return None

    cache_id = hash_archivos(rutas_contexto)
    carpeta_cache = RAG_DIR / cache_id
    carpeta_cache.mkdir(parents=True, exist_ok=True)

    try:
        embedding_function = SentenceTransformerEmbeddingFunction(
            model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
        )

        client = chromadb.PersistentClient(path=str(carpeta_cache))

        collection = client.get_or_create_collection(
            name=f"contexto_{cache_id[:20]}",
            embedding_function=embedding_function
        )

        try:
            cantidad = collection.count()
        except Exception:
            cantidad = 0

        if cantidad == 0:
            print("🔨 Construyendo RAG por primera vez...")

            chunks = dividir_en_chunks(texto_contexto)

            if not chunks:
                return None

            collection.add(
                documents=chunks,
                ids=[f"chunk_{i}" for i in range(len(chunks))],
                metadatas=[{"indice": i} for i in range(len(chunks))]
            )

        else:
            print("✅ RAG desde cache.")

        return collection

    except Exception as e:
        print(f"Error construyendo RAG: {e}")
        return None


def recuperar_contexto_pregunta(
    collection,
    pregunta: PreguntaExamen,
    n_resultados=2,
    limite_caracteres=2200
) -> str:
    if collection is None:
        return ""

    try:
        cantidad = collection.count()
        if cantidad <= 0:
            return ""

        n = min(n_resultados, cantidad)
        query = re.sub(r"\s+", " ", pregunta.bloque).strip()[:650]

        resultados = collection.query(
            query_texts=[query],
            n_results=n
        )

        documentos = resultados.get("documents", [[]])[0]

        contexto = []
        total = 0

        for doc in documentos:
            if not doc:
                continue

            if total + len(doc) > limite_caracteres:
                break

            contexto.append(doc)
            total += len(doc)

        return "\n\n--- FRAGMENTO RELEVANTE ---\n\n".join(contexto).strip()

    except Exception as e:
        print(f"Error recuperando contexto: {e}")
        return ""


def construir_contexto_por_pregunta(collection, preguntas: list) -> str:
    secciones = []

    for p in preguntas:
        contexto = recuperar_contexto_pregunta(collection, p)

        if not contexto:
            contexto = "No se encontró contexto suficiente."

        secciones.append(f"""
========================
PREGUNTA GLOBAL {p.id_global}
{p.serie} - Pregunta {p.numero_local}
VALOR MÁXIMO: {p.valor} puntos
========================

PREGUNTA Y RESPUESTA:
{p.bloque}

CONTEXTO RELEVANTE:
{contexto}
""")

    return "\n".join(secciones).strip()


# ============================================================
# PAQUETE DE EVALUACIÓN
# ============================================================

def preparar_paquete_evaluacion(rutas_imagenes: Union[str, Path, list], rutas_contexto: list) -> dict:
    """
    Función compatible con:
    - app viejo: preparar_paquete_evaluacion(ruta_imagen, rutas_contexto)
    - app nuevo: preparar_paquete_evaluacion([ruta1, ruta2, ruta3], rutas_contexto)
    """
    try:
        rutas_imagenes = normalizar_rutas(rutas_imagenes)
        rutas_contexto = normalizar_rutas(rutas_contexto)

        if not rutas_imagenes:
            return {"error": "No se recibió ninguna imagen del examen."}

        if not rutas_contexto:
            return {"error": "No se recibió ningún archivo de referencia."}

        imagen_hash = hash_archivos(rutas_imagenes)
        contexto_hash = hash_archivos(rutas_contexto)
        paquete_hash = hashlib.sha256(
            f"{imagen_hash}_{contexto_hash}".encode()
        ).hexdigest()

        cache_name = f"paquete_{paquete_hash}.json"
        cache = cargar_json_cache(cache_name)

        if cache:
            print("✅ Paquete desde cache.")
            preguntas = [PreguntaExamen(**p) for p in cache["preguntas"]]

            return {
                "texto_examen": cache.get("texto_examen", ""),
                "metadatos": cache.get("metadatos", {}),
                "rubrica": cache.get("rubrica", {}),
                "preguntas": preguntas,
                "tabla_rubrica": cache.get("tabla_rubrica", ""),
                "contexto_por_pregunta": cache.get("contexto_por_pregunta", ""),
                "imagenes_procesadas": cache.get("imagenes_procesadas", [Path(r).name for r in rutas_imagenes]),
            }

        texto_examen = extraer_texto_varias_imagenes(rutas_imagenes)

        if not texto_examen.strip():
            return {"error": "No se pudo extraer texto del examen usando Gemini."}

        if texto_examen.startswith("ERROR_IMAGEN_"):
            return {"error": texto_examen}

        if "NO HAY EXAMEN LEGIBLE" in texto_examen.upper():
            return {"error": "No se pudo leer un examen válido en una o varias imágenes."}

        metadatos = extraer_metadatos_examen(texto_examen)

        texto_contexto = leer_contextos(rutas_contexto)

        if not texto_contexto.strip():
            return {"error": "Los archivos de referencia no contienen texto útil o no se pudieron leer."}

        preguntas_raw = dividir_preguntas_con_series(texto_examen)

        if not preguntas_raw:
            return {"error": "No se detectaron preguntas en el examen."}

        rubrica = detectar_rubrica(texto_examen)
        preguntas = asignar_puntajes(preguntas_raw, rubrica)
        tabla_rubrica = construir_tabla_rubrica(preguntas, rubrica)

        collection = construir_rag_cache(rutas_contexto, texto_contexto)

        if collection is None:
            return {"error": "No se pudo construir o cargar el RAG del material de referencia."}

        contexto_por_pregunta = construir_contexto_por_pregunta(collection, preguntas)

        data_cache = {
            "texto_examen": texto_examen,
            "metadatos": metadatos,
            "rubrica": rubrica,
            "preguntas": [asdict(p) for p in preguntas],
            "tabla_rubrica": tabla_rubrica,
            "contexto_por_pregunta": contexto_por_pregunta,
            "imagenes_procesadas": [Path(r).name for r in rutas_imagenes],
        }

        guardar_json_cache(cache_name, data_cache)

        return {
            "texto_examen": texto_examen,
            "metadatos": metadatos,
            "rubrica": rubrica,
            "preguntas": preguntas,
            "tabla_rubrica": tabla_rubrica,
            "contexto_por_pregunta": contexto_por_pregunta,
            "imagenes_procesadas": [Path(r).name for r in rutas_imagenes],
        }

    except Exception as e:
        return {"error": f"Error preparando paquete de evaluación: {e}"}


# ============================================================
# DIFICULTAD
# ============================================================

def obtener_criterio_dificultad(nivel: int) -> str:
    return f"""
Nivel seleccionado: {nivel}/10

INTERPRETACIÓN GENERAL:
La dificultad controla qué tan completa, técnica y específica debe ser una respuesta para recibir el puntaje completo.
No significa castigar sin razón. Si el estudiante se acerca al concepto, debe recibir puntaje parcial proporcional.

ESCALA:
1 = Modo amigo extremo.
2 = Muy indulgente.
3 = Flexible.
4 = Moderadamente flexible.
5 = Estándar académico normal.
6 = Estándar exigente.
7 = Estricto.
8 = Muy estricto.
9 = Experto.
10 = Experto riguroso.

CRITERIO HUMANO DE CALIFICACIÓN:
- Respuesta totalmente correcta, técnica y completa: Correcta, 90% a 100%.
- Respuesta correcta pero breve: Parcial alta o Correcta según dificultad, 70% a 90%.
- Respuesta con idea central correcta pero incompleta: Parcial, 40% a 70%.
- Respuesta vaga, confusa o con errores importantes: Parcial baja, 10% a 40%.
- Respuesta totalmente incorrecta o en blanco: Incorrecta, 0%.

REGLAS POR NIVEL:
- Nivel 1-2: si la idea principal aparece, otorga casi todo el puntaje.
- Nivel 3-4: acepta respuestas breves si el concepto central es correcto.
- Nivel 5: acepta definiciones breves cuando la pregunta es de definición. Si la pregunta pide explicar, comparar, describir pasos o justificar, una respuesta muy breve debe ser Parcial.
- Nivel 6: exige concepto correcto y precisión moderada.
- Nivel 7: exige precisión y penaliza omisiones relevantes.
- Nivel 8: exige desarrollo suficiente. Respuestas generales en preguntas explicativas deben ser Parciales.
- Nivel 9: exige vocabulario técnico, completitud y explicación clara.
- Nivel 10: exige precisión técnica, completitud, relación con el material y claridad. No regales puntos completos por respuestas incompletas.

TOPES OBLIGATORIOS POR DIFICULTAD:
- Nivel 5:
  * Definición corta correcta: puede recibir 1.0 si cubre la idea esencial.
  * Pregunta explicativa muy breve: máximo 0.80.
  * Pregunta de pasos con pasos incompletos: máximo 0.70.
- Nivel 8:
  * Respuesta correcta pero muy breve en pregunta explicativa: máximo 0.75.
  * Pregunta que pide pasos y solo menciona 1 o 2 pasos generales: máximo 0.55.
  * Pregunta que pide comparar y solo da una diferencia superficial: máximo 0.65.
  * Definición correcta pero sin detalle técnico importante: máximo 0.85.
- Nivel 9:
  * Respuesta correcta pero breve en pregunta explicativa: máximo 0.65.
  * Pregunta de pasos incompleta: máximo 0.45.
  * Respuesta vaga con idea central: máximo 0.40.
- Nivel 10:
  * Respuesta completa, técnica y alineada al material: 0.90 a 1.00.
  * Respuesta correcta pero breve: 0.60 a 0.80.
  * Respuesta con idea central correcta pero incompleta: 0.30 a 0.60.
  * Respuesta vaga, confusa o con error importante: 0.10 a 0.30.
  * Respuesta sin relación o en blanco: 0.00.
  * Para dar 1.00 en nivel 10, la respuesta debe cubrir los elementos principales del material, no solo una frase general.

TIPOS DE PREGUNTA:
- Definición: pide definir un concepto.
- Explicativa: pide explicar, justificar, relacionar, describir, comparar, mencionar ventajas o dar pasos.
- Compuesta: pide varios elementos en una sola pregunta.

REGLAS PARA PREGUNTAS EXPLICATIVAS:
- Si pide "explique", debe haber explicación, no solo una frase.
- Si pide "diferencia entre A y B", debe mencionar características de ambos y la diferencia central.
- Si pide "pasos", debe listar o explicar la secuencia principal.
- Si pide "mencione y describa", no basta con mencionar; debe describir brevemente.
- Si pide ventajas, deben ser ventajas específicas, no frases vagas como "es mejor" o "da mejores respuestas".

IMPORTANTE:
Si el RAG no recupera una definición clara, usa conocimiento académico general siempre que no contradiga el material.
No califiques nivel {nivel} como si fuera otro nivel.
No uses "Correcta" solo porque la respuesta tiene una parte verdadera; si falta desarrollo relevante, usa "Parcial".
"""


# ============================================================
# CALIFICACIÓN
# ============================================================

def calificar_paquete(paquete: dict, nivel_dificultad: int) -> str:
    try:
        if not paquete:
            return "ERROR: No se recibió paquete de evaluación."

        if "error" in paquete:
            return f"ERROR: {paquete['error']}"

        rubrica = paquete.get("rubrica", {})
        texto_examen = paquete.get("texto_examen", "")
        tabla_rubrica = paquete.get("tabla_rubrica", "")
        contexto_por_pregunta = paquete.get("contexto_por_pregunta", "")
        metadatos = paquete.get("metadatos", {})
        criterio = obtener_criterio_dificultad(nivel_dificultad)
        total_rubrica = rubrica.get("total", 100)

        prompt = f"""
Eres un DOCENTE CALIFICADOR PROFESIONAL.
Califica este examen con precisión, justicia académica y criterio humano.

Usa exactamente:
1. El texto extraído del examen.
2. Los metadatos detectados.
3. La rúbrica calculada.
4. El contexto RAG por pregunta.
5. El nivel de dificultad indicado.

========================
METADATOS DETECTADOS
========================
{json.dumps(metadatos, ensure_ascii=False, indent=2)}

========================
EXAMEN EXTRAÍDO
========================
{texto_examen}

========================
RÚBRICA CALCULADA
========================
{tabla_rubrica}

========================
CONTEXTO POR PREGUNTA
========================
{contexto_por_pregunta}

========================
DIFICULTAD
========================
{criterio}

REGLAS OBLIGATORIAS:
- Usa exactamente el total de la rúbrica: {total_rubrica}.
- No inventes preguntas.
- No inventes respuestas.
- No cambies el valor máximo de cada pregunta.
- No redistribuyas puntos si ya hay valor por pregunta.
- No califiques sobre más de {total_rubrica}.
- Diferencia pregunta global y pregunta local.
- Al final convierte a escala de 100.
- La puntuación debe ser proporcional a cuánto se acerca la respuesta del estudiante a la respuesta esperada.
- No pongas 0 si el estudiante tiene parte de la idea correcta.
- No pongas 1.0 si la respuesta está incompleta para el nivel de dificultad seleccionado.
- Usa "Parcial" cuando la respuesta tenga una parte correcta pero le falte desarrollo, precisión, pasos, comparación o vocabulario técnico.
- Usa "Correcta" solamente cuando cubra la idea central y los elementos importantes esperados para ese nivel.
- Usa "Incorrecta" solamente cuando la respuesta sea falsa, no responda, esté en blanco o no tenga relación real con la pregunta.
- En preguntas de definición de 1 punto, una respuesta breve puede ser completa si contiene la esencia exacta del concepto.
- En preguntas explicativas de 1 punto, una respuesta muy breve no debe recibir 1.0 en niveles 5 o superiores.
- En nivel 10, una respuesta de una sola frase en preguntas explicativas casi siempre debe ser Parcial, no Correcta.
- Si el RAG trae una respuesta clara, úsala como base principal.
- Si el RAG no trae definición clara, usa conocimiento académico general si no contradice el material.
- Justifica cada puntuación con claridad.
- No seas excesivamente amable. No seas injustamente duro. Sé proporcional.

FORMATO FINAL OBLIGATORIO.
Responde SOLO en Markdown:

# Informe de Calificación

## Resumen General

**Punteo obtenido:** X/{total_rubrica}

**Nota final en escala de 100:** XX/100

**Nivel de dificultad aplicado:** {nivel_dificultad}/10

**Fuente de puntuación:** {rubrica.get("fuente", "No especificada")}

## Datos Detectados del Examen

**Estudiante:** ...

**Código o carné:** ...

**Curso:** ...

**Docente:** ...

**Título del examen:** ...

**Serie del examen:** ...

**Fecha del examen:** ...

## Detalle por Pregunta

### Pregunta global 1

**Serie:** ...

**Pregunta:** ...

**Respuesta del estudiante:** ...

**Respuesta esperada según el material:** ...

**Valor máximo:** ... puntos

**Puntuación obtenida:** ... puntos

**Estado:** Correcta / Parcial / Incorrecta

**Justificación:** ...

**Retroalimentación:** ...

Repite el bloque anterior para cada pregunta.

## Clasificación Final

**Buenas:** ...

**Parciales:** ...

**Incorrectas:** ...

## Conclusión General

...

## Recomendaciones de Estudio

...
"""

        reporte = llamar_gemini(prompt)

        if not reporte.strip():
            return "ERROR: Gemini terminó, pero no devolvió un reporte válido."

        return reporte

    except Exception as e:
        return f"ERROR: No se pudo calificar el paquete: {e}"


# ============================================================
# FUNCIÓN PARA APP.PY
# ============================================================

def calificar_examen_ui(rutas_imagenes: Union[str, Path, list], rutas_contexto: list, nivel_dificultad: int) -> str:
    paquete = preparar_paquete_evaluacion(rutas_imagenes, rutas_contexto)

    if "error" in paquete:
        return paquete["error"]

    return calificar_paquete(paquete, nivel_dificultad)
